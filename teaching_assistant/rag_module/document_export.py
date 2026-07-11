from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Import bộ lọc để đảm bảo nội dung sạch trước khi đưa vào Word
from .latex_formatter import process_science_formulas

def export_to_docx(content, filename="tai_lieu_giang_day.docx"):
    """
    Chuẩn hóa nội dung và xuất ra file Word (.docx) chuyên nghiệp.
    """
    doc = Document()
    
    # 1. Làm sạch nội dung qua bộ lọc chuẩn
    clean_content = process_science_formulas(content)
    
    # 2. Xử lý xóa bỏ các ký hiệu LaTeX đặc thù để Word hiển thị thuần văn bản
    # Chuyển đổi công thức frac{a}{b} thành a/b để Word đọc được
    def clean_latex(text):
        # Chuyển đổi \frac{a}{b} -> a/b
        text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1/\2', text)
        # Loại bỏ các ký hiệu LaTeX còn sót lại
        text = text.replace('$', '').replace('\\text', '').replace('{', '').replace('}', '')
        return text

    # 3. Tạo file Word
    lines = clean_content.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            # Nếu là tiêu đề (bắt đầu bằng # hoặc ##)
            if line.startswith('#'):
                header = doc.add_heading(level=1)
                header.text = line.replace('#', '').strip()
            else:
                # Nếu là nội dung thường
                paragraph = doc.add_paragraph(clean_latex(line))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                paragraph.style.font.size = Pt(12)
    
    doc.save(filename)
    return filename
