import streamlit as st
import sqlite3
import pandas as pd
import io

DB_PATH = "teacher_assistant.db"

def setup_minutes_database():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS org_minutes (
        id INTEGER PRIMARY KEY AUTOINCREMENT, meeting_date TEXT, session_number TEXT UNIQUE,
        present_members TEXT, absent_members TEXT, content_text TEXT, resolution TEXT
    )
    """)
    conn.commit()
    conn.close()

def extract_text_from_minutes_upload(uploaded_file):
    import docx
    from pypdf import PdfReader
    text_content = ""
    try:
        if uploaded_file.name.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            for p in doc.paragraphs:
                if p.text: text_content += p.text + "\n"
        elif uploaded_file.name.endswith('.pdf'):
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                text_content += (page.extract_text() or "") + "\n"
        elif uploaded_file.name.endswith('.txt'):
            text_content = uploaded_file.read().decode("utf-8")
    except Exception as e:
        st.error(f"Lỗi bóc tách file: {e}")
    return text_content.strip()

def export_minutes_to_docx(meeting_date, session_number, present_members, absent_members, content_text, resolution):
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n").bold = True
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(f"\nBIÊN BẢN SINH HOẠT TỔ CHUYÊN MÔN\n(Số: {session_number})\n")
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(255, 0, 0)
    items = [f"- Ngày họp: {meeting_date}", f"- Thành phần: {present_members}", f"- Vắng mặt: {absent_members}", "\n--- DIỄN BIẾN CUỘC HỌP ---", content_text, "\n--- NGHỊ QUYẾT CUỘC HỌP ---", resolution]
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(item)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
