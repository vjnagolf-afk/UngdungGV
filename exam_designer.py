# exam_designer.py
import streamlit as st
import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
import matplotlib.pyplot as plt
import numpy as np

def read_uploaded_docx(uploaded_file):
    try:
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    except: return "Lỗi đọc file Word"

def read_uploaded_pdf(uploaded_file):
    try:
        reader = PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    except: return "Lỗi đọc file PDF"

def generate_plot_stream(eq_str):
    fig, ax = plt.subplots(figsize=(5, 3.5))
    x = np.linspace(-10, 10, 400)
    safe_dict = {"x": x, "np": np, "sin": np.sin, "cos": np.cos, "tan": np.tan, "sqrt": np.sqrt}
    try:
        eq_str_py = eq_str.replace('^', '**')
        y = eval(eq_str_py, {"__builtins__": {}}, safe_dict)
        if isinstance(y, (int, float)):
            y = np.full_like(x, y)
        ax.plot(x, y, color='#1E40AF', linewidth=2.5)
        ax.axhline(0, color='black', linewidth=1.2)
        ax.axvline(0, color='black', linewidth=1.2)
        ax.grid(color='gray', linestyle='--', linewidth=0.5, alpha=0.7)
        ax.set_ylim([-10, 10])
        ax.set_title(f"Đồ thị: y = {eq_str}", fontsize=10, pad=10)
    except:
        ax.text(0.5, 0.5, f"[Không thể vẽ đồ thị: Sai cú pháp toán học]", ha='center', va='center', color='red')
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf

def export_to_docx_vietnam_standard(text_content, title_name, school_name="TRƯỜNG THCS NGUYỄN CHÍ THANH", group_name="TỔ KHOA HỌC TỰ NHIÊN - GDTC"):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    admin_table = doc.add_table(rows=1, cols=2)
    admin_table.autofit = False
    admin_table.columns[0].width = Inches(3.2)
    admin_table.columns[1].width = Inches(3.8)
    
    cell_l = admin_table.rows[0].cells[0]
    p_left = cell_l.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.add_run(f"{school_name.upper()}\n").bold = True
    p_left.add_run(f"{group_name.upper()}\n").bold = True
    p_left.add_run("Số: ..... /BB-TCM").font.size = Pt(11)
    
    cell_r = admin_table.rows[0].cells[1]
    p_right = cell_r.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n").bold = True
    p_right.add_run("Độc lập - Tự do - Hạnh phúc\n").bold = True
    p_right.add_run("***************").font.size = Pt(11)
    
    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run(title_name.upper()).bold = True
    
    in_table = False
    table_data = []
    
    def process_runs(paragraph, text):
        bold_parts = text.split('**')
        for i, b_part in enumerate(bold_parts):
            is_bold = (i % 2 != 0)
            sub_sup_parts = re.split(r'(<sub>.*?</sub>|<sup>.*?</sup>)', b_part)
            for part in sub_sup_parts:
                if not part: continue
                if part.startswith('<sub>') and part.endswith('</sub>'):
                    run = paragraph.add_run(part[5:-6]) 
                    run.bold = is_bold
                    run.font.subscript = True 
                elif part.startswith('<sup>') and part.endswith('</sup>'):
                    run = paragraph.add_run(part[5:-6]) 
                    run.bold = is_bold
                    run.font.superscript = True 
                else:
                    run = paragraph.add_run(part)
                    run.bold = is_bold

    def build_table():
        if not table_data: return
        cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_val in enumerate(row):
                if c_idx < cols:
                    cell = table.cell(r_idx, c_idx)
                    p = cell.paragraphs[0]
                    p.text = "" 
                    process_runs(p, cell_val.strip())
        doc.add_paragraph() 
        
    for line in text_content.split('\n'):
        cleaned_line = line.strip()
        if cleaned_line.startswith('|') and cleaned_line.endswith('|'):
            in_table = True
            row_data = [cell.strip() for cell in cleaned_line.split('|')[1:-1]]
            if all(re.match(r'^[-: ]+$', cell) for cell in row_data): continue
            table_data.append(row_data)
            continue
        if in_table:
            build_table()
            in_table = False
            table_data = []
        if not cleaned_line: continue
        if '[GRAPH:' in cleaned_line:
            match = re.search(r'\[GRAPH:\s*(.+?)\]', cleaned_line)
            if match:
                eq = match.group(1)
                img_stream = generate_plot_stream(eq)
                doc.add_picture(img_stream, width=Inches(3.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue 
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if cleaned_line.startswith('#'):
            process_runs(p, cleaned_line.replace('#', '').strip())
            for run in p.runs: run.bold = True
        else:
            process_runs(p, cleaned_line)
    if in_table: build_table()
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
def render_exam_designer_section(api_key_input, run_ai_prompt_safe_func):
    # CSS tạo khối tiêu đề màu Hồng (Phần Trắc nghiệm) và màu Xanh (Phần Tự luận) đúng hệt ảnh mẫu
    st.markdown("""
    <style>
    .section-pink { background-color: #FCE4EC; color: #880E4F; padding: 6px; text-align: center; font-weight: bold; font-size: 15px; border-radius: 4px; margin-bottom: 12px; letter-spacing: 1px;}
    .section-green { background-color: #E8F5E9; color: #1B5E20; padding: 6px; text-align: center; font-weight: bold; font-size: 15px; border-radius: 4px; margin-bottom: 12px; letter-spacing: 1px;}
    .footer-red { color: #D32F2F; font-weight: bold; font-style: italic; font-size: 14px; text-align: center; margin-top: 30px; padding-top: 10px; border-top: 1px solid #ccc;}
    div[data-testid="stForm"] { border: none !important; padding: 0 !important; }
    </style>
    """, unsafe_allow_html=True)

    if "db_de_kiem_tra" not in st.session_state: st.session_state["db_de_kiem_tra"] = []
    if "current_exam_designer_output" not in st.session_state: st.session_state["current_exam_designer_output"] = ""

    tab_thiet_ke, tab_kho_luu_tru = st.tabs(["📝 CHỨC NĂNG: TẠO ĐỀ KIỂM TRA AI", "📂 THƯ MỤC ĐỀ ĐÃ XÂY DỰNG"])
    
    with tab_thiet_ke:
        # Hàng nút chức năng trên cùng phía bên phải
        col_top_lbl, col_top_btn1, col_top_btn2 = st.columns([3, 1.2, 1.2])
        hinh_thuc = col_top_lbl.selectbox("Hình thức đề:", ["Trắc nghiệm kết hợp tự luận", "100% Trắc nghiệm", "100% Tự luận"], label_visibility="visible")
        col_top_btn1.markdown("<div style='margin-top:28px;'></div>", unsafe_allow_html=True)
        col_top_btn1.button("TÁI KHUNG MA TRẬN MẪU", type="secondary", use_container_width=True)
        col_top_btn2.markdown("<div style='margin-top:28px;'></div>", unsafe_allow_html=True)
        tai_lieu_dinh_kem = col_top_btn2.file_uploader("TÁI TÀI LIỆU TẢI LÊN", type=["docx", "pdf"], label_visibility="collapsed")
        
        st.markdown("<p style='color: gray; font-size: 12px; font-style: italic; margin-top:-10px;'>ℹ️ Chưa có tài liệu nào được tải lên hệ thống.</p>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)

        # Bố cục 2 Cột chính: PHẦN TRẮC NGHIỆM (Hồng) & PHẦN TỰ LUẬN (Xanh) đúng theo ảnh của thầy
        col_main1, col_main2 = st.columns(2)
        
        # --- 💥 CỘT LỆCH TRÁI: PHẦN TRẮC NGHIỆM ---
        with col_main1:
            st.markdown("<div class='section-pink'>PHẦN TRẮC NGHIỆM</div>", unsafe_allow_html=True)
            
            # Khởi tạo các ô nhập số câu và điểm thành phần của trắc nghiệm
            r_tn1, r_tn2 = st.columns(2)
            c1_tn = r_tn1.number_input("Số câu nhiều lựa chọn:", min_value=0, max_value=50, value=12, step=1, key="num_tn_1")
            p1_tn = r_tn2.number_input("Tổng điểm dòng này (Nhiều lựa chọn):", min_value=0.0, max_value=10.0, value=2.0, step=0.25, key="pt_tn_1")
            
            r_tn3, r_tn4 = st.columns(2)
            c2_tn = r_tn3.number_input("Số câu đúng sai:", min_value=0, max_value=50, value=2, step=1, key="num_tn_2")
            p2_tn = r_tn4.number_input("Tổng điểm dòng này (Đúng sai):", min_value=0.0, max_value=10.0, value=1.0, step=0.25, key="pt_tn_2")
            
            r_tn5, r_tn6 = st.columns(2)
            c3_tn = r_tn5.number_input("Số câu điền khuyết:", min_value=0, max_value=50, value=1, step=1, key="num_tn_3")
            p3_tn = r_tn6.number_input("Tổng điểm dòng này (Điền khuyết):", min_value=0.0, max_value=10.0, value=0.5, step=0.25, key="pt_tn_3")
            
            r_tn7, r_tn8 = st.columns(2)
            c4_tn = r_tn7.number_input("Số câu trả lời ngắn:", min_value=0, max_value=50, value=1, step=1, key="num_tn_4")
            p4_tn = r_tn8.number_input("Tổng điểm dòng này (Trả lời ngắn):", min_value=0.0, max_value=10.0, value=0.5, step=0.25, key="pt_tn_4")
            
            # Logic Cộng dồn tự động: Tính tổng số câu hỏi TN và tổng số điểm TN
            tong_cau_tn_calc = c1_tn + c2_tn + c3_tn + c4_tn
            tong_diem_tn_calc = p1_tn + p2_tn + p3_tn + p4_tn
            
            # Đè hộp hiển thị tổng số câu hỏi lên phía đầu phần trắc nghiệm
            st.markdown(
                f"<div style='display:flex; justify-content:space-between; margin-top:-340px; margin-bottom:310px;'>"
                f"<div style='font-weight:bold; color:#9F1239;'>Tổng số câu TNKQ: <span style='background:white; padding:4px 20px; border:1px solid #D1D5DB; border-radius:4px; margin-left:10px;'>{tong_cau_tn_calc}</span></div>"
                f"<div style='font-weight:bold; color:#9F1239;'>Tổng điểm TN: <span style='background:white; padding:4px 20px; border:1px solid #D1D5DB; border-radius:4px; margin-left:10px;'>{tong_diem_tn_calc:.1f}</span></div>"
                f"</div>", unsafe_allow_html=True
            )

        # --- 💥 CỘT LỆCH PHẢI: PHẦN TỰ LUẬN (TỰ ĐỘNG SINH CÂU VÀ CỘNG DỒN) ---
        with col_main2:
            st.markdown("<div class='section-green'>PHẦN TỰ LUẬN</div>", unsafe_allow_html=True)
            
            # Ô cấu hình tổng số câu tự luận
            num_tl_input = st.number_input("TỔNG SỐ CÂU TỰ LUẬN:", min_value=0, max_value=20, value=5, step=1, key="num_tl_root_v8")
            
            # Thuật toán vòng lặp tự sinh các dòng Câu con và ô điểm tương ứng đúng theo ảnh mẫu
            tl_scores = []
            if num_tl_input > 0:
                for i in range(num_tl_input):
                    r_cell1, r_cell2, r_cell3 = st.columns([1, 3, 1])
                    r_cell1.markdown(f"<div style='margin-top:8px;'>Câu {i+1}</div>", unsafe_allow_html=True)
                    score_val = r_cell2.number_input(f"Score_TL_{i}", min_value=0.0, max_value=10.0, value=1.0 if i!=1 and i!=2 else 1.5, step=0.5, label_visibility="collapsed")
                    r_cell3.markdown("<div style='margin-top:8px; font-size:12px; color:gray;'>ĐIỂM</div>", unsafe_allow_html=True)
                    tl_scores.append(score_val)
            
            # Tính tổng số điểm tự luận từ tổng các câu con
            tong_diem_tl_calc = sum(tl_scores)
            
            # Đè hộp hiển thị tổng số câu và tổng điểm tự luận lên phía trên giống hệt ảnh mẫu của thầy
            st.markdown(
                f"<div style='display:flex; justify-content:space-between; margin-top:-40px; margin-bottom:20px; position: relative; top: -{num_tl_input*44 + 50}px;'>"
                f"<div style='font-weight:bold; color:#1E3A8A; position:absolute; top:-43px; left:0px;'>TỔNG SỐ CÂU TỰ LUẬN: <span style='background:white; padding:4px 15px; border:1px solid #D1D5DB; border-radius:4px; margin-left:5px;'>{num_tl_input}</span></div>"
                f"<div style='font-weight:bold; color:#1E3A8A; position:absolute; top:-43px; right:0px;'>ĐIỂM: <span style='background:white; padding:4px 15px; border:1px solid #D1D5DB; border-radius:4px; margin-left:5px;'>{tong_diem_tl_calc:.1f}</span></div>"
                f"</div>", unsafe_allow_html=True
            )
            # Tạo khoảng trống bù lề để đẩy form tiếp theo xuống dưới không bị đè chữ
            st.markdown(f"<div style='margin-top:-30px;'></div>", unsafe_allow_html=True)

        st.markdown("---")
        # Khu vực hàng nút lệnh Khởi tạo AI và Checkbox bám sát tài liệu
        col_btn_zone, col_check_zone = st.columns([1.5, 2])
        run_exam_ai = col_btn_zone.button("⚙️ Tự động tạo ma trận & đề thi", type="primary", use_container_width=True, key="btn_run_exam_ultimate")
        yeu_cau_bam_sat = col_check_zone.checkbox("Yêu cầu bám sát kiến thức trong tài liệu tải lên", value=True, key="chk_bam_sat_ultimate")

        # Cấu hình 4 ô nhập tỷ lệ mức độ nhận thức nằm ngang hàng
        st.markdown("**Tỷ lệ mức độ nhận thức (%):**")
        col_mz1, col_mz2, col_mz3, col_mz4 = st.columns(4)
        mz_nb = col_mz1.number_input("Nhận biết:", min_value=0, max_value=100, value=40, step=5, key="mz_nb_u")
        mz_th = col_mz2.number_input("Thông hiểu:", min_value=0, max_value=100, value=30, step=5, key="mz_th_u")
        mz_vd = col_mz3.number_input("Vận dụng:", min_value=0, max_value=100, value=20, step=5, key="mz_vd_u")
        mz_vdc = col_mz4.number_input("Vận dụng cao:", min_value=0, max_value=100, value=10, step=5, key="mz_vdc_u")

        # Ô nhập yêu cầu bổ sung dạng văn bản diện rộng
        note_de = st.text_area("Nhập yêu cầu khác ....", placeholder="Nhập yêu cầu khác ....", label_visibility="collapsed", key="note_de_area_u")

        # --- LOGIC XỬ LÝ GỌI AI CHẠY NGẦM BẰNG API KEY HỆ THỐNG ---
        if run_exam_ai:
            context_text = ""
            if tai_lieu_dinh_kem is not None:
                if tai_lieu_dinh_kem.name.endswith(".docx"): context_text = read_uploaded_docx(tai_lieu_dinh_kem)
                elif tai_lieu_dinh_kem.name.endswith(".pdf"): context_text = read_uploaded_pdf(tai_lieu_dinh_kem)

            prompt_exam = (
                f"Hãy thiết kế một Ma trận đề thi và Đề kiểm tra chi tiết (kèm Đáp án) cho môn: Khoa học tự nhiên, hình thức: {hinh_thuc}, thời gian: 45 phút.\n"
                f"CẤU TRÚC ĐỀ BẮT BUỘC ĐỒNG BỘ:\n"
                f"- Trắc nghiệm tổng cộng {tong_cau_tn_calc} câu với {tong_diem_tn_calc} điểm, chi tiết gồm: Many-choice: {c1_tn} câu, True/False: {c2_tn} câu, Fill-blank: {c3_tn} câu, Short-answer: {c4_tn} câu.\n"
                f"- Tự luận tổng cộng {num_tl_input} câu lớn với tổng điểm là {tong_diem_tl_calc} điểm (Điểm chi tiết từng câu con: {', '.join([f'Câu {i+1}: {s}đ' for i, s in enumerate(tl_scores)])}).\n"
                f"MA TRẬN MỨC ĐỘ NHẬN THỨC: Nhận biết {mz_nb}%, Thông hiểu {mz_th}%, Vận dụng {mz_vd}%, Vận dụng cao {mz_vdc}%.\n"
                f"Yêu cầu nội dung kiến thức đính kèm: {note_de}.\n"
                f"Tài liệu tham khảo đính kèm:\n{context_text}\n\n"
                f"Trả về văn bản định dạng Markdown đẹp mắt, ma trận dựng bảng dạng '|'."
            )

            if run_ai_prompt_safe_func is not None:
                with st.spinner("🚀 Trợ lý AI đang liên kết API Key hệ thống và thiết kế đề thi..."):
                    res_text, status = run_ai_prompt_safe_func(prompt_exam)
                    if status == "error":
                        st.error(f"❌ Lỗi kết nối: {res_text}")
                    else:
                        st.session_state["current_exam_designer_output"] = res_text
                        st.session_state["db_de_kiem_tra"].append({
                            "id": f"De_{len(st.session_state['db_de_kiem_tra'])+1}", "mon": "Khoa học tự nhiên", "khoi": "Khối 7", "hinh_thuc": hinh_thuc, "data": res_text
                        })
                        st.success(f"🎉 Khởi tạo đề kiểm tra thành công bằng mô hình {status}!")
                        st.rerun()
            else:
                st.error("❌ Lỗi luồng: Chưa kết nối được trình điều khiển AI tổng từ file app.py.")

        # Khung hiển thị văn bản đầu ra và nút Tải file Word chuẩn văn bản Việt Nam
        with st.container(border=True):
            if st.session_state["current_exam_designer_output"]:
                st.markdown(st.session_state["current_exam_designer_output"])
                
                word_exam_data = export_to_docx_vietnam_standard(st.session_state["current_exam_designer_output"], "ĐỀ KIỂM TRA MÔN KHOA HỌC TỰ NHIÊN")
                st.download_button(label="🍏 Tải về file Word (.docx)", data=word_exam_data, file_name="De_Kiem_Tra_KHTN.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            else:
                st.caption("Đề thi và ma trận xuất hiện ở đây")

    # --- KHU VỰC THƯ MỤC ĐỀ ĐÃ XÂY DỰNG (LỊCH SỬ KHO LƯU TRỮ) ---
    with tab_kho_luu_tru:
        st.markdown("#### 📂 Các đề kiểm tra đã được AI tự động sinh và lưu trữ")
        if not st.session_state["db_de_kiem_tra"]:
            st.caption("✨ Chưa có đề kiểm tra nào được tạo.")
        else:
            indices_to_delete = []
            for idx, de_info in enumerate(st.session_state["db_de_kiem_tra"]):
                with st.expander(f"📋 Đề môn {de_info['mon']} - {de_info['hinh_thuc']}"):
                    st.markdown(de_info['data'])
                    h_col1, h_col2 = st.columns(2)
                    word_hist_data = export_to_docx_vietnam_standard(de_info['data'], f"ĐỀ KIỂM TRA MÔN {de_info['mon'].upper()}")
                    h_col1.download_button(label="📥 Tải file Word", data=word_hist_data, file_name=f"De_{idx+1}.docx", key=f"dl_ex_w_{idx}", use_container_width=True)
                    if h_col2.button("❌ Xóa đề thi", key=f"del_ex_btn_{idx}", use_container_width=True): indices_to_delete.append(idx)
            if indices_to_delete:
                for index in sorted(indices_to_delete, reverse=True): st.session_state["db_de_kiem_tra"].pop(index)
                st.success("🗑️ Đã xóa đề thi!")
                st.rerun()

    st.markdown("<div class='footer-red'>© Bản quyền thuộc về Tác giả: Lê Hồng Dưỡng | Đơn vị: Trường THCS Nguyễn Chí Thanh - phường Tân Lập - tỉnh Đắk Lắk</div>", unsafe_allow_html=True)
