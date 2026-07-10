import streamlit as st
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime

# Cấu hình Sheets
SHEET_ID = '1C6642jk_oQ0g9UC2By2qsNxxfQVR0MrZYj52tRdWDlY'

def get_danhgia_sheet():
    try:
        creds_dict = dict(st.secrets["GOOGLE_KEY"])
        scopes = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
        creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)
        client = gspread.authorize(creds)
        return client.open_by_key(SHEET_ID).worksheet("DANH_GIA_HS")
    except Exception as e:
        st.error(f"Lỗi kết nối Google Sheets: {e}")
        return None

def set_paragraph_spacing(paragraph, before_pt=3.0, after_pt=4.5):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    p_pr.append(spacing)

def create_word_table(doc, table_data, MAU_DEN):
    """Hàm bổ trợ để vẽ bảng tránh lặp code"""
    if not table_data:
        return
    num_rows = len(table_data)
    num_cols = len(table_data[0]) if num_rows > 0 else 0
    
    if num_cols > 0:
        word_table = doc.add_table(rows=num_rows, cols=num_cols)
        word_table.style = 'Table Grid'
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                if c_idx < num_cols:
                    cell = word_table.cell(r_idx, c_idx)
                    cell.text = val
                    for para in cell.paragraphs:
                        set_paragraph_spacing(para, 2.0, 3.0)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for r in para.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(12) # 12pt phù hợp cho bảng hơn 14pt để tránh tràn dòng
                            r.font.color.rgb = MAU_DEN

def export_rubric_to_docx(title_text, markdown_content):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)

    MAU_DO = RGBColor(255, 0, 0)
    MAU_XANH_DUONG = RGBColor(0, 51, 153)
    MAU_DEN = RGBColor(0, 0, 0)

    # Tiêu đề Rubric căn giữa màu đỏ
    p_title = doc.add_paragraph()
    set_paragraph_spacing(p_title, 6.0, 6.0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(title_text.upper())
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = MAU_DO

    lines = markdown_content.split('\n')
    in_table = False
    table_data = []

    for line in lines:
        clean_line = line.strip().replace('**', '').replace('###', '').replace('##', '').replace('#', '')
        
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if '---|' in line or ':---|' in line: 
                continue
            in_table = True
            cells = [c.strip().replace('**', '') for c in line.split('|')[1:-1]]
            table_data.append(cells)
            continue
        else:
            if in_table and table_data:
                create_word_table(doc, table_data, MAU_DEN)
                in_table = False
                table_data = []

        if not clean_line: 
            continue

        if re.match(r'^(I|II|III|IV|V)\.', clean_line) or re.match(r'^\d+\.', clean_line):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, 4.0, 4.5)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(clean_line)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = MAU_XANH_DUONG
        else:
            p = doc.add_paragraph()
            set_paragraph_spacing(p)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(clean_line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    # Xử lý nếu bảng nằm ở cuối cùng văn bản văn bản mẫu
    if in_table and table_data:
        create_word_table(doc, table_data, MAU_DEN)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

import streamlit as st
# Đảm bảo đã cài đặt: pip install google-genai
from google import genai
from google.genai import types

def call_gemini_with_fallback(prompt_text, preferred_model):
    """
    Hàm gọi API Gemini có cơ chế tự động chuyển đổi mô hình (Fallback)
    """
    # Khởi tạo client sử dụng API Key từ Streamlit secrets
    try:
        client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    except Exception as e:
        return f"❌ Lỗi cấu hình API Key: {str(e)}", None

    # Định nghĩa danh sách thứ tự mô hình ưu tiên hạ cấp (Fallback)
    # Tên mô hình chính thức cập nhật theo tài liệu Google GenAI
    model_pool = {
        "3.1 Pro": ["gemini-1.5-pro", "gemini-1.5-flash", "gemini-2.5-flash-8b"],
        "3.5 Flash": ["gemini-1.5-flash", "gemini-2.5-flash-8b"],
        "Flash Lite": ["gemini-2.5-flash-8b"]
    }
    
    # Lấy danh sách chuỗi mô hình sẽ thử nghiệm dựa trên lựa chọn của người dùng
    models_to_try = model_pool.get(preferred_model, ["gemini-1.5-flash"])
    
    last_error = ""
    for model_name in models_to_try:
        try:
            # Thông báo trạng thái đang thử nghiệm mô hình nào
            # (Hữu ích khi debug, có thể ẩn đi nếu muốn giao diện sạch hơn)
            response = client.models.generate_content(
                model=model_name,
                contents=prompt_text,
                config=types.GenerateContentConfig(
                    temperature=0.3, # Thấp một chút để Rubric chuẩn xác, định lượng tốt
                )
            )
            # Nếu thành công, trả về kết quả và tên mô hình đã chạy thành công
            return response.text, model_name
        except Exception as e:
            last_error = str(e)
            # Nếu lỗi, vòng lặp for sẽ tự động chuyển sang model_name tiếp theo trong danh sách
            continue
            
    # Nếu thử toàn bộ danh sách vẫn lỗi thì trả về thông báo lỗi cuối cùng
    return f"❌ Tất cả các mô hình đều gặp lỗi. Lỗi cuối cùng: {last_error}", None


def render_assessment_section(): # Bỏ tham số run_ai_prompt_safe_func vì đã tích hợp trực tiếp bên trên
    st.markdown("<h3 style='text-align: center; color: red;'>🎯 TRỢ LÝ THIẾT KẾ RUBRIC ĐÁNH GIÁ HỌC SINH TỰ ĐỘNG</h3>", unsafe_allow_html=True)
    
    tab_thiet_ke, tab_thu_vien = st.tabs(["📝 THIẾT KẾ TIÊU CHÍ RUBRIC AI", "🗄️ LƯU TRỮ RUBRIC"])
    
    if "ket_qua_rubric" not in st.session_state: 
        st.session_state["ket_qua_rubric"] = ""
    if "model_da_dung" not in st.session_state:
        st.session_state["model_da_dung"] = ""

    with tab_thiet_ke:
        st.write("Nhập tên nội dung bài học hoặc chủ đề để hệ thống tự động thiết kế bảng tiêu chí đánh giá định lượng.")
        
        noi_dung = st.text_input("Tên nội dung kiến thức bài học / Chương / Chủ đề / Sản phẩm:", placeholder="Ví dụ: Mô hình xe phản lực - Chương Tốc độ chuyển động")
        
        col_lop, col_loai, col_model = st.columns([1, 1.5, 1.5])
        with col_lop:
            lop = st.text_input("Lớp:", placeholder="Ví dụ: 7A")
        with col_loai:
            hinh_thuc = st.selectbox("Hình thức đánh giá:", ["Qua sản phẩm học tập", "Qua bài thuyết trình", "Qua hoạt động nhóm", "Đánh giá năng lực thực hành"])
        with col_model:
            # 🆕 THÊM HỘP CHỌN MÔ HÌNH ƯU TIÊN
            mo_hinh_uu_tien = st.selectbox("Mô hình ưu tiên sử dụng:", ["3.1 Pro", "3.5 Flash", "Flash Lite"])

        col_btn1, col_blank, col_btn2 = st.columns([2.5, 1.0, 2.0])
        
        with col_btn2:
            st.write(""); st.write("")
            nut_rubric = st.button("⚡ Khởi tạo tiêu chí bằng AI", type="primary", use_container_width=True)

        if nut_rubric:
            if not noi_dung:
                st.warning("⚠️ Vui lòng điền nội dung kiến thức bài học hoặc tên sản phẩm!")
            else:
                with st.spinner("🧠 Hệ thống đang phân tích mục tiêu bài học để lập bảng Rubric..."):
                    prompt_rubric = f"""
                    Bạn là Chuyên gia kiểm tra đánh giá giáo dục phổ thông tại Việt Nam. Hãy thiết kế một bảng tiêu chí đánh giá định lượng (Rubric) chi tiết cho:
                    - Nội dung/Sản phẩm: {noi_dung}
                    - Lớp: {lop}
                    - Hình thức đánh giá: {hinh_thuc}
                    
                    TIÊU ĐỀ KHỐI (Viết in hoa ở dòng đầu):
                    BẢNG TIÊU CHÍ RUBRIC ĐÁNH GIÁ: {noi_dung.upper()} - LỚP {lop.upper()}
                    
                    YÊU CẦU CẤU TRÚC VÀ ĐỊNH DẠNG VĂN BẢN:
                    1. BẢNG TIÊU CHÍ RUBRIC: Phải xây dựng dưới dạng bảng Markdown hoàn chỉnh bằng ký tự '|'. Các cột bao gồm: Tiêu chí đánh giá, Trọng số (%), Mức Tốt (8.0-10 điểm), Mức Khá (6.5-7.9 điểm), Mức Đạt (5.0-6.4 điểm), Mức Chưa đạt (Dưới 5.0 điểm).
                    2. NỘI DUNG TIÊU CHÍ: Phải bám sát đặc thù kiến thức bài học/sản phẩm của học sinh lớp {lop}. Mô tả rõ ràng hành vi định lượng để giáo viên dễ chấm, không viết chung chung.
                    3. BỎ TOÀN BỘ các ký tự dấu sao kép '**' ở đầu và cuối các từ hoặc tiêu đề cột để tránh lỗi font chữ đậm.
                    4. Danh sách liệt kê dùng dấu gạch ngang '-' ở đầu dòng.
                    """
                    # 🆕 GỌI HÀM CÓ CHỨA CƠ CHẾ TỰ ĐỘNG CHUYỂN ĐỔI MÔ HÌNH
                    ket_qua, model_thuc_te = call_gemini_with_fallback(prompt_rubric, mo_hinh_uu_tien)
                    st.session_state["ket_qua_rubric"] = ket_qua
                    st.session_state["model_da_dung"] = model_thuc_te

        # Hiển thị kết quả AI và thông báo rõ mô hình nào đã phản hồi thành công
        if st.session_state["ket_qua_rubric"]:
            if st.session_state["model_da_dung"]:
                st.info(f"🤖 Đã tạo thành công bằng mô hình: `{st.session_state['model_da_dung']}`")
            st.markdown("### 📋 Kết quả Rubric từ AI:")
            st.markdown(st.session_state["ket_qua_rubric"])

        with col_btn1:
            st.write(""); st.write("")
            title_file = f"Rubric_{lop}_{noi_dung[:20].replace(' ', '_')}" if noi_dung else "Rubric_Danh_Gia"
            
            if st.session_state["ket_qua_rubric"] and "❌ Lỗi" not in st.session_state["ket_qua_rubric"]:
                from danh_gia_manager import export_rubric_to_docx # Hoặc import từ file của bạn
                docx_data = export_rubric_to_docx(f"BẢNG RUBRIC ĐÁNH GIÁ: {noi_dung}", st.session_state["ket_qua_rubric"])
                is_disabled = False
            else:
                docx_data = b""
                is_disabled = True

            st.download_button(
                label="📥 Tải file Word (.docx) chuẩn về máy",
                data=docx_data,
                file_name=f"{title_file}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=is_disabled,
                use_container_width=True
            )

