import streamlit as st  # <-- ĐÃ VÁ LỖI CÚ PHÁP THÊM CHỮ "AS" TẠI ĐÂY
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- HÀM XUẤT PHÔI WORD KHĐ THEO PHỤ LỤC III BỘ GIÁO DỤC ---
def export_plan_to_docx(teacher_name, subject_name, markdown_content):
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n").bold = True
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("\nKẾ HOẠCH GIÁO DỤC CỦA GIÁO VIÊN\n(PHỤ LỤC III CÔNG VĂN 5512/BGDĐT)\n")
    r_title.bold = True
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = RGBColor(255, 0, 0)
    p_info = doc.add_paragraph()
    p_info.add_run(f"Giáo viên thực hiện: {teacher_name}\nMôn học/Hoạt động giáo dục: {subject_name}\n")
    lines = markdown_content.split('\n')
    for line in lines:
        clean_line = line.strip().replace('**', '').replace('###', '').replace('##', '').replace('#', '')
        if not clean_line: continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(clean_line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        if line.strip().startswith('#') or "I." in clean_line or "II." in clean_line:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def render_personal_plan():
    st.markdown("<h3 style='text-align: left; color: #1E3A8A;'>🗓️ TRỢ LÝ XÂY DỰNG KẾ HOẠCH GIÁO DỤC CÁ NHÂN (PHỤ LỤC III)</h3>", unsafe_allow_html=True)
    
    # --- KHU VỰC PHÂN QUYỀN MẬT KHẨU ADMIN TẠI SIDEBAR ---
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🔒 CHỌN VAI TRÒ ĐĂNG NHẬP")
    vai_tro = st.sidebar.radio("Vai trò", ["Giáo viên bộ môn", "Tổ trưởng chuyên môn (Admin)", "Ban giám hiệu"], label_visibility="collapsed", key="vai_tro_plan_sidebar_fixed_v8")
    
    # Giữ lại xác thực mã PIN cho vai trò Tổ trưởng/BGH nếu các tính năng khác trong hệ sinh thái của bạn cần dùng
    if vai_tro == "Tổ trưởng chuyên môn (Admin)":
        ma_pin = st.sidebar.text_input("Nhập mã pin quản lý Admin:", type="password", value="", key="pin_admin_plan_v8")
        if ma_pin == "123456":
            st.sidebar.success("✅ Quyền Admin đã mở.")
        elif ma_pin != "": 
            st.sidebar.error("❌ Mã PIN sai.")
            return

    if "ai_plan_output" not in st.session_state:
        st.session_state["ai_plan_output"] = ""

    # Hộp chứa giao diện tĩnh đứng im trên màn hình chống lag widget
    container_output = st.empty()
    
    with st.form("form_personal_plan_fixed_final_v8", border=False):
        col_t, col_s = st.columns(2)
        t_name = col_t.text_input("Họ và tên Giáo viên giảng dạy:", placeholder="Ví dụ: Thầy Lê Hồng Dưỡng", key="plan_txt_t_name_v8")
        s_name = col_s.selectbox("Môn học / Phân môn phụ trách:", ["Khoa học tự nhiên (Vật lý)", "Khoa học tự nhiên (Sinh học)", "Khoa học tự nhiên (Hóa học)", "Toán học", "Ngữ văn", "GDTC"], key="plan_sb_s_name_v8")
        col_g, col_w = st.columns(2)
        grade_target = col_g.text_input("Khối lớp phân công dạy:", placeholder="Ví dụ: Khối 7, Khối 8", key="plan_txt_grade_target_v8")
        week_count = col_w.text_input("Tổng số tuần thực hiện kế hoạch:", placeholder="Ví dụ: 35 Tuần", key="plan_txt_week_count_v8")
        st.markdown("**💬 Các tiêu chí đặc thù hoặc lưu ý phân bổ tiết (Nếu có):**")
        note_plan = st.text_area("Yêu cầu bổ sung cho AI:", placeholder="Ví dụ: Phân bổ chi tiết số tiết cho chương Tốc độ ở vật lý 7 học kỳ I...", label_visibility="collapsed", key="plan_ta_note_v8")
        
        # Nút bấm khởi tạo được mở cho toàn bộ giáo viên
        run_ai_plan = st.form_submit_button("🚀 Khởi tạo Kế hoạch bằng AI", type="primary", use_container_width=True)
        
    # --- ĐOẠN CODE ĐÃ ĐƯỢC MỞ KHÓA CHO TẤT CẢ GIÁO VIÊN ---
    if run_ai_plan:
        if not t_name or not grade_target:
            st.warning("⚠️ Vui lòng điền Họ tên giáo viên và Khối lớp để AI lập kế hoạch!")
        else:
            ai_success = False
            with st.spinner("Trợ lý AI đang lập tiến trình phân bổ tiết Phụ lục III..."):
                prompt_plan = f"Hãy soạn thảo một bản Kế hoạch giáo dục của giáo viên (Phụ lục III - Công văn 5512) chi tiết cho giáo viên: {t_name}, môn: {s_name}, khối lớp: {grade_target} trong {week_count}. Cấu trúc gồm mục I. KẾ HOẠCH DẠY HỌC PHÂN PHỐI CHƯƠNG TRÌNH và mục II. CÁC NHIỆM VỤ KHÁC ĐƯỢC GIAO. Văn bản viết chi tiết đầy đủ chữ, chia dòng gạch ngang '-', không dùng dấu sao kép '**'."
                try:
                    from app import run_ai_prompt_safe
                    api_key_system = st.secrets.get("GEMINI_API_KEY", "")
                    
                    if not api_key_system:
                        st.error("🔑 Thiếu cấu hình GEMINI_API_KEY trong Secrets!")
                        st.stop()

                    res_plan, status = run_ai_prompt_safe(prompt_plan, api_key_system)
                    st.session_state["ai_plan_output"] = res_plan
                    ai_success = True
                    
                except Exception as ai_error:
                    st.error(f"❌ Máy chủ AI hoặc mã nguồn kết nối phản hồi lỗi: {str(ai_error)}")
                    st.info("Vui lòng kiểm tra lại hạn ngạch API Gemini hoặc cấu hình mạng máy chủ.")
            
            if ai_success:
                st.rerun()

    st.markdown("---")
    st.markdown("### 📊 Nội dung Kế hoạch Giáo dục sinh bởi AI:")
    with st.container(border=True):
        if st.session_state["ai_plan_output"]:
            st.markdown(st.session_state["ai_plan_output"])
            word_plan_data = export_plan_to_docx(t_name, s_name, st.session_state["ai_plan_output"])
            st.download_button(label="📥 Tải file Word (.docx) Phụ lục III chuẩn về máy", data=word_plan_data, file_name=f"Phu_Luc_III_{t_name.replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, key="download_fixed_v8")
        else:
            st.caption("Khung kế hoạch chi tiết sẽ xuất hiện tại đây sau khi Tổ trưởng hoặc Giáo viên bấm nút khởi tạo...")
