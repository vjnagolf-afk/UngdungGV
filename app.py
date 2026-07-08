# app.py
import streamlit as st
import io
import json
import os
import pandas as pd
import time
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from pypdf import PdfReader

# Gọi hàm thiết kế đề thi và quản lý điểm từ file module đã tách
from exam_designer import render_exam_designer_section
from grade_manager import render_grade_manager_section
from timetable_manager import render_timetable_section
# ==============================================================================
# 1. CẤU HÌNH TRANG VÀ THIẾT LẬP SIÊU CSS ĐỊNH DẠNG THANH BÊN THEO BIỂU MẪU CHUẨN
# ==============================================================================
st.set_page_config(
    page_title="HỆ SINH THÁI SỐ GIÁO VIÊN",
    page_icon="🚀",
    layout="wide"
)

st.title("📚 HỆ SINH THÁI SỐ - HỖ TRỢ GIÁO VIÊN")
st.caption("Sản phẩm tham gia Cuộc thi AI for Life năm 2026, trường THCS Nguyễn Chí Thanh - Phường Tân Lập tỉnh Đắk Lắk")
st.markdown("---")

st.markdown("""
<style>
div[data-testid="stSidebarUserContent"] { padding-top: 1.0rem !important; padding-bottom: 0px !important; }
div[data-testid="stSidebarUserContent"] hr { margin-top: 0.6rem !important; margin-bottom: 0.6rem !important; }
div[data-testid="stSidebarUserContent"] div[data-testid="stVerticalBlock"] > div { margin-bottom: 0px !important; }
div[data-testid="stSidebarUserContent"] div[data-testid="stRadio"] > label {
    font-weight: 900 !important; color: #E11D48 !important; text-align: center !important;
    display: block !important; width: 100% !important; font-size: 26px !important;
    letter-spacing: 0.5px !important; border-bottom: 2px dashed #FDA4AF !important;
    padding-bottom: 6px !important; margin-bottom: 12px !important;
}
div[data-testid="stSidebarUserContent"] div[data-testid="stRadio"] div[data-baseweb="radio"] div[data-testid="stMarkdownContainer"] p { font-weight: 900 !important; font-size: 22px !important; line-height: 1.4 !important; }
div[data-testid="stSidebarUserContent"] div[data-testid="stRadio"] div[data-baseweb="radio"]:nth-child(1) div[data-testid="stMarkdownContainer"] p { color: #1D4ED8 !important; }
div[data-testid="stSidebarUserContent"] div[data-testid="stRadio"] div[data-baseweb="radio"]:nth-child(2) div[data-testid="stMarkdownContainer"] p { color: #EA580C !important; }
div.stButton > button, div.stDownloadButton > button { transition: all 0.3s cubic-bezier(0.25, 0.8, 0.25, 1) !important; border-radius: 8px !important; }
div.stButton > button:hover, div.stDownloadButton > button:hover { background-color: #1E40AF !important; color: #FFFFFF !important; border-color: #3B82F6 !important; transform: translateY(-2px) !important; box-shadow: 0 6px 15px rgba(59, 130, 246, 0.4) !important; }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 2. BỘ ENGINE AI & KẾ XUẤT VĂN BẢN CHUẨN ĐỊNH DẠNG HÀNH CHÍNH (TIMES NEW ROMAN)
# ==============================================================================
def run_ai_prompt_safe(prompt_text, api_key):
    models_to_try = ["gemini-2.5-pro", "gemini-2.5-flash", "gemini-1.5-flash"]
    client = genai.Client(api_key=api_key)
    last_error = None
    for model in range(len(models_to_try)):
        for attempt in range(2):
            try:
                response = client.models.generate_content(model=models_to_try[model], contents=prompt_text)
                return response.text, models_to_try[model]
            except Exception as e:
                last_error = e
                time.sleep(2)
    raise Exception(f"Tất cả các mô hình AI đều đang gặp lỗi: {last_error}")

def export_to_docx_vietnam_standard(text_content, title_name, school_name="TRƯỜNG THCS NGUYỄN CHÍ THANH", group_name="TỔ KHOA HỌC TỰ NHIÊN - GDTC"):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.59)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    admin_table = doc.add_table(rows=1, cols=2)
    admin_table.autofit = False
    admin_table.columns[0].width = Inches(3.2)
    admin_table.columns[1].width = Inches(3.8)
    
    cell_l = admin_table.rows[0].cells[0]
    p_left = cell_l.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.add_run(f"{school_name.upper()}\n").bold = True
    p_left.add_run(f"{group_name.upper()}\n").bold = True
    p_left.add_run("Số: ..... /BB-TCM").font.size = Pt(11)
    
    cell_r = admin_table.rows[0].cells[1]
    p_right = cell_r.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n").bold = True
    p_right.add_run("Độc lập - Tự do - Hạnh phúc\n").bold = True
    p_right.add_run("***************").font.size = Pt(11)
    
    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run(title_name.upper()).bold = True
    
    in_table = False
    table_data = []
    
    def build_table():
        if not table_data: return
        cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, cell_val in enumerate(row):
                if c_idx < cols:
                    clean_val = cell_val.replace('**', '').strip()
                    table.cell(r_idx, c_idx).text = clean_val
        doc.add_paragraph() 
        
    for line in text_content.split('\n'):
        cleaned_line = line.strip()
        
        if cleaned_line.startswith('|') and cleaned_line.endswith('|'):
            in_table = True
            row_data = [cell.strip() for cell in cleaned_line.split('|')[1:-1]]
            if all(re.match(r'^[-: ]+$', cell) for cell in row_data):
                continue
            table_data.append(row_data)
            continue
            
        if in_table:
            build_table()
            in_table = False
            table_data = []
            
        if not cleaned_line: continue
            
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if cleaned_line.startswith('#'):
            p.add_run(cleaned_line.replace('#', '').strip()).bold = True
        elif '**' in cleaned_line:
            parts = cleaned_line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True
        else:
            p.add_run(cleaned_line)
            
    if in_table:
        build_table()
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def convert_df_to_excel_bytes(df, sheet_name='Sheet1'):
    val_output = io.BytesIO()
    with pd.ExcelWriter(val_output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return val_output.getvalue()

def generate_dynamic_5512_docx():
    doc = Document()
    title = doc.add_paragraph()
    title.add_run("KHUNG KẾ HOẠCH BÀI DẠY MẪU CHUẨN 5512").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def read_uploaded_docx(uploaded_file):
    try:
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    except: return "Lỗi đọc file Word"

def read_uploaded_pdf(uploaded_file):
    try:
        reader = PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    except: return "Lỗi đọc file PDF"

# ==============================================================================
# HÀM PHỤ TRỢ: LƯU TRỮ VÀ ĐỌC HỒ SƠ 
# ==============================================================================
def save_data_to_local_file(file_name, data):
    try:
        with open(file_name, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
    except: pass

def load_data_from_local_file(file_name, default_data):
    if os.path.exists(file_name):
        try:
            with open(file_name, "r", encoding="utf-8") as f:
                return json.load(f)
        except: return default_data
    return default_data

FILE_THANH_VIEN = "db_thanh_vien.json"
FILE_THANH_TICH = "db_thanh_tich_da_nam.json"
DANH_SACH_NAM_HOC = [f"{nam} - {nam+1}" for nam in range(2020, 2036)]

MAU_THANH_VIEN = [
    {"STT": 1, "Họ và tên": "Lê Hồng Dưỡng", "Chức vụ": "Tổ trưởng", "Phân môn chính": "KHTN (Lý) - CN", "Email": "vjnagolf@gmail.com", "Số điện thoại": "0984331178"},
    {"STT": 2, "Họ và tên": "Nguyễn Thị Huyền Trang", "Chức vụ": "Tổ viên", "Phân môn chính": "KHTN (Lý) - CN", "Email": "binhnt@gmail.com", "Số điện thoại": "0987654321"},
    {"STT": 3, "Họ và tên": "Lê Hùng Cường", "Chức vụ": "Tổ viên", "Phân môn chính": "KHTN (Lý) - CN", "Email": "viettt@gmail.com", "Số điện thoại": "0905123456"}
]

MAU_THANH_TICH = {
    "2025 - 2026": [
        {"Năm học": "2025 - 2026", "Họ và tên": "Lê Hồng Dưỡng", "Đánh giá viên chức": "Hoàn thành xuất sắc (HTSX)", "Kết quả BD HSG": "1 giải Nhất môn Vật lý", "Kết quả NCKH": "Đạt giải Nhì", "Kết quả STEM": "02 bài học xuất sắc", "Kết quả Sáng kiến": "Giải B", "Thi GVDG": "Cấp Tỉnh", "Thi GVCNG": "Không", "Kết quả TDTT": "Giải Nhất bóng bàn", "Kết quả VN": "Giải xuất sắc", "Hiến máu nhân đạo": "01 lần", "Hoạt động khác": "Tích cực Chuyển đổi số"}
    ]
}

if "db_thanh_vien" not in st.session_state: st.session_state["db_thanh_vien"] = load_data_from_local_file(FILE_THANH_VIEN, MAU_THANH_VIEN)
if "db_thanh_tich_da_nam" not in st.session_state: st.session_state["db_thanh_tich_da_nam"] = load_data_from_local_file(FILE_THANH_TICH, MAU_THANH_TICH)

def sync_sub_databases():
    current_pc = st.session_state.get("db_phan_cong_hien_tai", [])
    pc_mapped = {x["Họ và tên"]: x for x in current_pc}
    new_pc = []
    
    for idx, tv in enumerate(st.session_state["db_thanh_vien"]):
        name = tv.get("Họ và tên", "Giáo viên mới")
        try: stt_val = int(tv["STT"]) if (pd.notna(tv.get("STT")) and str(tv.get("STT")).strip() != "") else (idx + 1)
        except: stt_val = idx + 1
            
        if name in pc_mapped:
            old_item = pc_mapped[name]
            new_pc.append({
                "STT": stt_val, "Họ và tên": name, "Phân môn chính": tv.get("Phân môn chính", ""),
                "Lớp phân công": old_item.get("Lớp phân công", "9A1, 9A2"), "Số tiết phân công": old_item.get("Số tiết phân công", 8),
                "Nhiệm vụ kiêm nhiệm": tv.get("Chức vụ", "")
            })
        else:
            new_pc.append({
                "STT": stt_val, "Họ và tên": name, "Phân môn chính": tv.get("Phân môn chính", ""),
                "Lớp phân công": "", "Số tiết phân công": 0, "Nhiệm vụ kiêm nhiệm": tv.get("Chức vụ", "")
            })
    st.session_state["db_phan_cong_hien_tai"] = new_pc

if "db_phan_cong_hien_tai" not in st.session_state: sync_sub_databases()
if "db_de_kiem_tra" not in st.session_state:
    st.session_state["db_de_kiem_tra"] = [
        {"ten_de": "Đề giữa kỳ I - KHTN 9", "mon": "Khoa học tự nhiên", "khoi": "Lớp 9", "noi_dung": "### MA TRẬN ĐỀ THI ĐÃ DỰNG"}
    ]

has_secrets = "GEMINI_API_KEY" in st.secrets if hasattr(st, "secrets") else False
api_key_input = st.secrets["GEMINI_API_KEY"] if has_secrets else ""
st.sidebar.title("MENU HỆ THỐNG")
phan_he_lam_viec = st.sidebar.radio("CHỌN PHÂN HỆ TÁC NGHIỆP", [" Trợ lý Giảng dạy (Giáo viên)", " Trợ lý Quản lý (Tổ chuyên môn)"], index=0)

if phan_he_lam_viec == " Trợ lý Giảng dạy (Giáo viên)":
    chuc_nang_chinh = st.sidebar.selectbox("CHỌN NỘI DUNG THỰC HIỆN", ["1. Thiết kế KHBD thông minh", "2. Thiết kế Đề KTĐG (Ma trận - Đặc tả)", "3. Đánh giá học sinh","4. Quản lý điểm học sinh (SMAS)","5. Quản lý Thời khóa biểu"])
else:
    chuc_nang_chinh = st.sidebar.selectbox("QUẢN LÝ TỔ CHUYÊN MÔN", ["1. Hệ thống Quản lý và Phân công chuyên môn giảng dạy", "2. Xây dựng Biên bản sinh hoạt tổ chuyên môn định kỳ", "3. Xây dựng Kế hoạch Giáo dục cá nhân (Phụ lục III - Công văn 5512)", "4. Thống kê số liệu tổ"])

st.sidebar.markdown("---")
st.sidebar.title("🔒 CHỌN VAI TRÒ ĐĂNG NHẬP")
user_role = st.sidebar.selectbox("VỚI TƯ CÁCH LÀ:", ["Giáo viên / tổ viên", "Tổ trưởng chuyên môn (Admin)"], index=1)
is_admin_choice = (user_role == "Tổ trưởng chuyên môn (Admin)")

if "is_admin_verified" not in st.session_state: st.session_state["is_admin_verified"] = False

if is_admin_choice:
    admin_password = st.sidebar.text_input("Nhập mã pin quản trị Admin:", type="password", key="pwd_admin_root")
    if admin_password == "021476":
        st.session_state["is_admin_verified"] = True
        st.sidebar.success("✅ Xác thực thành công! Quyền Admin đã mở.")
    else:
        st.session_state["is_admin_verified"] = False
        if admin_password: st.sidebar.error("❌ Sai mã pin bảo mật!")
        else: st.sidebar.info("💡 Điền mã pin để kích hoạt quyền sửa.")
else:
    st.session_state["is_admin_verified"] = False
    st.sidebar.info("💡 Bạn được xem hồ sơ tổ")

if not has_secrets: api_key_input = st.sidebar.text_input("Nhập khóa Gemini API Key nếu cần dùng AI:", type="password", value=api_key_input)
st.sidebar.markdown("---")
tac_gia = st.sidebar.text_input("Họ và tên tác giả:", value="Lê Hồng Dưỡng")
don_vi = st.sidebar.text_input("Đơn vị công tác:", value="Trường THCS Nguyễn Chí Thanh")

# ==============================================================================
# GIAO DIỆN CHÍNH
# ==============================================================================
if phan_he_lam_viec == " Trợ lý Giảng dạy (Giáo viên)":
    if chuc_nang_chinh == "1. Thiết kế KHBD thông minh":
        st.header("📋 THIẾT KẾ KẾ HOẠCH BÀI DẠY (KHBD) TÍCH HỢP NĂNG LỰC SỐ")
        st.info("Thầy cô có thể tải trực tiếp tài liệu hướng dẫn cấu trúc khung kế hoạch bài dạy mẫu 5512 thế hệ mới:")
        col_w, col_t = st.columns(2)
        with col_w: st.download_button(label="📥 Tải Khung KHBD chuẩn (.docx)", data=generate_dynamic_5512_docx(), file_name="Khung_Ke_hoach_Bai_day_5512.docx")
        with col_t: st.download_button(label="📥 Tải Cẩm nang sư phạm (.txt)", data="CẨM NANG SƯ PHẠM 5512 NÂNG CAO".encode('utf-8'), file_name="Cam_nang_Huong_dan_5512.txt")
        
        st.markdown("---")
        col_input1, col_input2 = st.columns(2)
        with col_input1:
            ten_bai = st.text_input("Tên bài học học tập:", placeholder="Ví dụ: Bài 1: Qua Đèo Ngang")
            lop = st.selectbox("Khối lớp giảng dạy:", ["Lớp 6", "Lớp 7", "Lớp 8", "Lớp 9"], index=3)
        with col_input2:
            mon_hoc = st.text_input("Môn học chuyên môn:", value="Khoa học tự nhiên")
            thoi_luong = st.text_input("Thời lượng thực hiện:", value="1 tiết")
            
        uploaded_khbd_file = st.file_uploader("Tải file học liệu gốc/Nội dung SGK nguồn (.pdf, .docx):", type=["pdf", "docx"], key="khbd_up")
        khbd_content = ""
        if uploaded_khbd_file:
            khbd_content = read_uploaded_docx(uploaded_khbd_file) if uploaded_khbd_file.name.endswith('.docx') else read_uploaded_pdf(uploaded_khbd_file)
            st.success("✅ Đã nạp thành công học liệu gốc thành công vào bộ nhớ AI.")
            
        uu_tien_tai_lieu = st.checkbox("🎯 ƯU TIÊN BÁM SÁT 100% TÀI LIỆU TẢI LÊN", value=True)
        if st.button("🚀 XD KHBD thông minh"):
            with st.spinner("AI đang dựng giáo án..."):
                try:
                    res, _ = run_ai_prompt_safe(f"Soạn giáo án chuẩn 5512 cho bài {ten_bai}", api_key_input)
                    st.markdown(res)
                except Exception as e: st.error(f"Lỗi: {e}")
                
    elif chuc_nang_chinh == "2. Thiết kế Đề KTĐG (Ma trận - Đặc tả)":
        # TÍCH HỢP TỪ MODULE ĐÃ TÁCH
        render_exam_designer_section(api_key_input, run_ai_prompt_safe)

    elif chuc_nang_chinh == "3. Đánh giá học sinh":
        st.header("📋 BỘ CÔNG CỤ ĐÁNH GIÁ, NHẬN XÉT VÀ HỖ TRỢ HỌC SINH")
        chu_de = st.text_input("Nội dung đánh giá:", value="Thầy/Cô điền tên bài học/Chủ đề cần xây dựng tiêu chí đánh giá vào đây")
        if st.button("🚀 Sinh học liệu"):
            with st.spinner("AI đang tạo..."):
                try:
                    res, _ = run_ai_prompt_safe(f"Tạo tiêu chí rubric cho chủ đề {chu_de}", api_key_input)
                    st.markdown(res)
                except Exception as error_rb: st.error(f"Lỗi: {error_rb}")

    elif chuc_nang_chinh == "4. Quản lý điểm học sinh (SMAS)":
        render_grade_manager_section() # GỌI HÀM VẼ GIAO DIỆN QUẢN LÝ ĐIỂM

else:
    # PHÂN HỆ QUẢN LÝ TỔ CHUYÊN MÔN
    if chuc_nang_chinh == "1. Hệ thống Quản lý và Phân công chuyên môn giảng dạy":
        st.subheader("📋 HỆ THỐNG QUẢN LÝ VÀ PHÂN CÔNG CHUYÊN MÔN GIẢNG DẠY")
        tab_thanh_vien, tab_phan_cong, tab_thanh_tich = st.tabs(["👥 DANH SÁCH THÀNH VIÊN & CHUYÊN MÔN", "📅 THEO DÕI PHÂN CÔNG CHUYÊN MÔN", "📈 GHI CHÉP THÀNH TÍCH CỦA GIÁO VIÊN"])
        
        with tab_thanh_vien:
            df_thanh_vien = pd.DataFrame(st.session_state["db_thanh_vien"])
            df_tv_edited = st.data_editor(df_thanh_vien, num_rows="dynamic" if st.session_state["is_admin_verified"] else "fixed", disabled=not st.session_state["is_admin_verified"], use_container_width=True)
            if st.session_state["is_admin_verified"] and st.button("💾 Lưu cập nhật Danh sách"):
                st.session_state["db_thanh_vien"] = df_tv_edited.to_dict('records')
                save_data_to_local_file(FILE_THANH_VIEN, st.session_state["db_thanh_vien"])
                sync_sub_databases() 
                st.success("✅ Đã cập nhật!")
                st.rerun()

        with tab_phan_cong:
            df_pc_current = pd.DataFrame(st.session_state["db_phan_cong_hien_tai"])
            df_pc_edited = st.data_editor(df_pc_current, num_rows="fixed", disabled=not st.session_state["is_admin_verified"], use_container_width=True)
            if st.session_state["is_admin_verified"] and st.button("💾 Lưu Phân công"):
                st.session_state["db_phan_cong_hien_tai"] = df_pc_edited.to_dict('records')
                st.success("✅ Đã lưu!")
                st.rerun()

        with tab_thanh_tich:
            nam_hoc_selected = st.selectbox("📅 CHỌN NĂM HỌC QUẢN LÝ:", DANH_SACH_NAM_HOC, index=5)
            if nam_hoc_selected not in st.session_state["db_thanh_tich_da_nam"] or not st.session_state["db_thanh_tich_da_nam"][nam_hoc_selected]:
                st.session_state["db_thanh_tich_da_nam"][nam_hoc_selected] = [{"Năm học": nam_hoc_selected, "Họ và tên": tv["Họ và tên"], "Đánh giá viên chức": "Hoàn thành tốt (HTT)", "Kết quả BD HSG": "Không", "Kết quả NCKH": "Không", "Kết quả STEM": "Không", "Kết quả Sáng kiến": "Không", "Thi GVDG": "Không", "Thi GVCNG": "Không", "Kết quả TDTT": "Không", "Kết quả VN": "Không", "Hiến máu nhân đạo": "Không", "Hoạt động khác": "Không"} for tv in st.session_state["db_thanh_vien"]]
            
            list_tt_data = st.session_state["db_thanh_tich_da_nam"][nam_hoc_selected]
            for row in list_tt_data: row["Năm học"] = nam_hoc_selected
            df_tt_current = pd.DataFrame(list_tt_data)
            df_tt_edited = st.data_editor(df_tt_current, num_rows="dynamic" if st.session_state["is_admin_verified"] else "fixed", disabled=not st.session_state["is_admin_verified"], use_container_width=True)
            
            if st.session_state["is_admin_verified"]:
                col_save_tt, col_ex_tt = st.columns(2)
                with col_save_tt:
                    if st.button(f"💾 Lưu dữ liệu năm {nam_hoc_selected}"):
                        saved_records = df_tt_edited.to_dict('records')
                        for record in saved_records: record["Năm học"] = nam_hoc_selected
                        st.session_state["db_thanh_tich_da_nam"][nam_hoc_selected] = saved_records
                        save_data_to_local_file(FILE_THANH_TICH, st.session_state["db_thanh_tich_da_nam"]) 
                        st.success("✅ Đã lưu!")
                        st.rerun()
                with col_ex_tt:
                    st.download_button(label=f"📥 Tải Bảng thi đua (.xlsx)", data=convert_df_to_excel_bytes(df_tt_current), file_name=f"Thi_Dua_{nam_hoc_selected.replace(' - ', '_')}.xlsx")

    elif chuc_nang_chinh == "2. Xây dựng Biên bản sinh hoạt tổ chuyên môn định kỳ":
        st.subheader("📝 XÂY DỰNG BIÊN BẢN SINH HOẠT TỔ CHUYÊN MÔN")
        st.markdown("#### I. THỜI GIAN, ĐỊA ĐIỂM, THÀNH PHẦN HOẠT ĐỘNG")
        col_t1, col_t2, col_t3 = st.columns(3)
        with col_t1:
            ky_hop_so = st.text_input("Kỳ họp thứ:", value="05")
            thang_hop_so = st.text_input("Biên bản họp tháng:", value="01")
            nam_hoc_chuoi = st.text_input("Năm học trường đăng ký:", value="2025 - 2026")
        with col_t2:
            gio_phut_hop = st.text_input("Vào lúc mấy giờ mấy phút:", value="13 giờ 30 phút")
            ngay_thang_hop = st.text_input("Ngày/Tháng/Năm cuộc họp diễn ra:", value="17/01/2026")
            phong_hop = st.text_input("Địa điểm tại phòng học:", value="Hội trường")
        with col_t3:
            chu_tri_cuoc = st.text_input("Chủ trì cuộc họp (Chủ tọa):", value="Thầy Lê Hồng Dưỡng – Tổ trưởng")
            thu_ky_cuoc = st.text_input("Thư ký lập biên bản:", value="Cô Nguyễn Thị Bình")
            thanh_phan_co = st.text_input("Thành phần (Có mặt / Vắng mặt):", value="Có mặt: 10/10; Vắng mặt: 0")

        st.markdown("---")
        st.markdown("#### II. NỘI DUNG CUỘC HỌP & NẠP KẾ HOẠCH NGUỒN")
        uploaded_kh_thang = st.file_uploader("📥 Tải file Kế hoạch tháng làm căn cứ biên bản (.pdf, .docx):", type=["pdf", "docx"], key="up_kh_thang")
        content_kh_thang = ""
        if uploaded_kh_thang:
            content_kh_thang = read_uploaded_docx(uploaded_kh_thang) if uploaded_kh_thang.name.endswith('.docx') else read_uploaded_pdf(uploaded_kh_thang)
            st.success("✅ Đã nạp thành công kế hoạch tháng vào bộ nhớ trợ lý ảo.")

        col_m1, col_m2, col_m3 = st.columns(3)
        with col_m1:
            bo_sung_11 = st.text_area("Ghi chú bổ sung mục 1.1 (Công tác trọng tâm):", placeholder="Nhập thêm lưu ý nếu có...")
        with col_m2:
            bo_sung_12 = st.text_area("Ghi chú bổ sung mục 1.2 (Thảo luận chuyên môn):", placeholder="Nhập lưu ý phương pháp, chuyên đề STEM...")
        with col_m3:
            bo_sung_13 = st.text_area("Ghi chú bổ sung mục 1.3 (Công tác khác):", placeholder="Nhập việc phong trào, công đoàn...")

        y_kien_dong_gop = st.text_area("Mục 2. Ý kiến đóng góp của các thành viên tổ chuyên môn:", value="- Các thành viên thảo luận sôi nổi và hoàn toàn nhất trí với phương hướng phân bổ chuyên môn tháng này.")

        if st.button("🤖 AI Tự động soạn thảo Biên bản"):
            if not uploaded_kh_thang: st.warning("⚠️ Thầy cần đính kèm file Kế hoạch tháng lên trước!")
            else:
                with st.spinner("AI đang dựng biên bản hành chính..."):
                    try:
                        prompt_bb = f"Hãy đóng vai thư ký tổ chuyên môn trường THCS Nguyễn Chí Thanh. Dựa vào nội dung file kế hoạch: {content_kh_thang}, soạn thảo nội dung phần II. NỘI DUNG CUỘC HỌP bóc tách lắp vào mục 1.1, 1.2, 1.3 phát triển chi tiết."
                        res_ai, _ = run_ai_prompt_safe(prompt_bb, api_key_input)
                        full_bien_ban_text = f"BIÊN BẢN SINH HOẠT TỔ CHUYÊN MÔN\nKỲ HỌP THỨ {ky_hop_so} THÁNG {thang_hop_so}\n\nI. THỜI GIAN, ĐỊA ĐIỂM\nLúc {gio_phut_hop}, ngày {ngay_thang_hop} tại {phong_hop}\nChủ trì: {chu_tri_cuoc}, Thư ký: {thu_ky_cuoc}, Thành phần: {thanh_phan_co}\n\nII. NỘI DUNG CUỘC HỌP\n{res_ai}\n\n2. Ý KIẾN ĐÓNG GÓP\n{y_kien_dong_gop}"
                        st.markdown(full_bien_ban_text)
                        st.download_button("📥 Xuất văn bản biên bản Word (.docx)", data=export_to_docx_vietnam_standard(full_bien_ban_text, f"Bien_ban_thang_{thang_hop_so}"), file_name=f"Bien_ban_thang_{thang_hop_so}.docx")
                    except Exception as e: st.error(f"Lỗi: {e}")

    elif chuc_nang_chinh == "3. Xây dựng Kế hoạch Giáo dục cá nhân (Phụ lục III - Công văn 5512)":
        st.header("📋 KẾ HOẠCH GIÁO DỤC CÁ NHÂN CỦA GIÁO VIÊN (Phụ lục III)")
        gv_selected = st.selectbox("Chọn Giáo viên cần lập kế hoạch:", [x["Họ và tên"] for x in st.session_state["db_thanh_vien"]])
        mon_tu_dong = next((x["Phân môn chính"] for x in st.session_state["db_thanh_vien"] if x["Họ và tên"] == gv_selected), "")
        st.info(f"Phân môn giảng dạy: **{mon_tu_dong}**")
        if st.button("🪄 AI Thiết kế Phụ lục III chuẩn 5512"):
            with st.spinner("AI đang xử lý..."):
                try:
                    res, _ = run_ai_prompt_safe(f"Thiết kế khung kế hoạch giáo dục cá nhân phụ lục III 5512 cho môn {mon_tu_dong}", api_key_input)
                    st.markdown(res)
                    st.download_button("📥 Tải Phụ lục III Word (.docx)", data=export_to_docx_vietnam_standard(res, f"Phu_luc_III_{gv_selected}"), file_name=f"Phu_luc_III_{gv_selected}.docx")
                except Exception as e: st.error(f"Lỗi: {e}")

   elif chuc_nang_chinh == "4. Thống kê số liệu tổ":
        st.header("📊 THỐNG KÊ GIÁO VIÊN TỔ")
        tong_nhan_su = len(st.session_state["db_thanh_vien"])
        st.metric("Tổng số thành viên tổ", f"{tong_nhan_su} Giáo viên")
        df_tv_current = pd.DataFrame(st.session_state["db_thanh_vien"])
        if "Phân môn chính" in df_tv_current.columns:
            st.bar_chart(df_tv_current["Phân môn chính"].value_counts())

    elif chuc_nang_chinh == "5. Quản lý Thời khóa biểu":
        render_timetable_section()  # Dòng này phải thụt vào trong như trên
