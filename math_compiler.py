# math_compiler.py - Bản Nâng Cấp Toàn Diện Hỗ Trợ Đầy Đủ Toán/Lý/Hóa
import io
import re
import numpy as np
import matplotlib.pyplot as plt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt

# ================= TỪ ĐIỂN KÝ HIỆU HY LẠP VÀ TOÁN HỌC =================
GREEK = {
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ', r'\epsilon': 'ε', r'\varepsilon': 'ε',
    r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ', r'\vartheta': 'ϑ', r'\iota': 'ι', r'\kappa': 'κ',
    r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π',
    r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ', r'\phi': 'φ', r'\varphi': 'ϕ',
    r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
    r'\Delta': 'Δ', r'\Gamma': 'Γ', r'\Theta': 'Θ', r'\Lambda': 'Λ', r'\Pi': 'Π',
    r'\Sigma': 'Σ', r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω'
}

SYMBOLS = {
    r'\infty': '∞', r'\rightarrow': '→', r'\leftarrow': '←', r'\Rightarrow': '⇒',
    r'\Leftarrow': '⇐', r'\leftrightarrow': '↔', r'\Leftrightarrow': '⇔',
    r'\approx': '≈', r'\neq': '≠', r'\leq': '≤', r'\geq': '≥', r'\times': '×',
    r'\cdot': '·', r'\pm': '±', r'\mp': '∓', r'\circ': '°', r'\partial': '∂',
    r'\nabla': '∇', r'\forall': '∀', r'\exists': '∃', r'\in': '∈', r'\notin': '∉',
    r'\subset': '⊂', r'\supset': '⊃', r'\cup': '∪', r'\cap': '∩', r'\equiv': '≡',
    r'\sim': '∼', r'\propto': '∝', r'\angle': '∠', r'\parallel': '∥', r'\perp': '⊥',
    r'\int': '∫', r'\sum': '∑'
}

# Ánh xạ marker Unicode trung gian sang các thẻ OMML
TAG_MAP = {
    'Ⓕ': '<m:f>', 'ⓕ': '</m:f>',
    'Ⓝ': '<m:num>', 'ⓝ': '</m:num>',
    'Ⓓ': '<m:den>', 'ⓓ': '</m:den>',
    'Ⓠ': '<m:rad><m:radPr><m:deg m:val=""/></m:radPr>', 'ⓠ': '</m:rad>',
    'Ⓔ': '<m:e>', 'ⓔ': '</m:e>',
    'Ⓑ': '<m:sSub>', 'ⓑ': '</m:sSub>',
    'Ⓟ': '<m:sSup>', 'ⓟ': '</m:sSup>',
    'Ⓩ': '<m:sSubSup>', 'ⓩ': '</m:sSubSup>',
    'Ⓢ': '<m:sub>', 'ⓢ': '</m:sub>',
    'Ⓤ': '<m:sup>', 'ⓤ': '</m:sup>',
    'Ⓧ': '<m:nary><m:naryPr><m:chr m:val="∫"/><m:limLoc m:val="subSup"/></m:naryPr>', 'ⓧ': '</m:nary>',
    'Ⓨ': '<m:nary><m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/></m:naryPr>', 'ⓨ': '</m:nary>',
    'Ⓥ': '<m:acc><m:accPr><m:chr m:val="&#x20D7;"/></m:accPr>', 'ⓥ': '</m:acc>',
    'Ⓜ': '<m:m><m:mPr><m:mcs><m:mc><m:mcPr><m:count m:val="10"/><m:mcJc m:val="c"/></m:mcPr></m:mc></m:mcs></m:mPr>', 'ⓜ': '</m:m>',
    'Ⓡ': '<m:mr>', 'ⓡ': '</m:mr>'
}

# ================= BIÊN DỊCH TOÁN HỌC LATEX SANG OMML WORD =================
def convert_latex_to_omml(latex_str):
    # 1. Làm sạch và mã hóa ký tự HTML
    latex_str = latex_str.strip()
    latex_str = latex_str.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    
    # 2. Xử lý các ký hiệu Hy Lạp, Toán học và loại bỏ các thẻ \text thừa
    for k, v in GREEK.items(): latex_str = latex_str.replace(k, v)
    for k, v in SYMBOLS.items(): latex_str = latex_str.replace(k, v)
    latex_str = latex_str.replace(r'\,', ' ').replace(r'\;', ' ')
    latex_str = re.sub(r'\\text\s*\{([^{}]*)\}', r'\1', latex_str)
    latex_str = re.sub(r'\\mathrm\s*\{([^{}]*)\}', r'\1', latex_str)

    # 3. Xử lý Ma trận (matrix, pmatrix...)
    def repl_matrix(m):
        rows = m.group(1).split('\\\\')
        res = 'Ⓜ'
        for row in rows:
            res += 'Ⓡ'
            cols = row.split('&')
            for col in cols: res += f'Ⓔ{col.strip()}ⓔ'
            res += 'ⓡ'
        res += 'ⓜ'
        return res
    latex_str = re.sub(r'\\begin\{(?:p|b|v|V)?matrix\}(.*?)\\end\{(?:p|b|v|V)?matrix\}', repl_matrix, latex_str, flags=re.DOTALL)

    # 4. Vòng lặp biên dịch lồng ghép (Từ trong ra ngoài)
    while True:
        prev = latex_str
        
        # Phân số (\frac)
        latex_str = re.sub(r'\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'ⒻⓃ\1ⓝⒹ\2ⓓⓕ', latex_str)
        # Căn bậc hai (\sqrt)
        latex_str = re.sub(r'\\sqrt\s*\{([^{}]+)\}', r'ⓆⒺ\1ⓔⓠ', latex_str)
        # Vector (\vec)
        latex_str = re.sub(r'\\vec\s*\{([^{}]+)\}', r'ⓋⒺ\1ⓔⓥ', latex_str)
        
        # Tích phân/Tổng lồng cả sub và sup (∫_a^b)
        latex_str = re.sub(r'(∫|∑)_\{([^{}]+)\}\^\{([^{}]+)\}', r'ⓍⒺ\1ⓔⓈ\2ⓢⓊ\3ⓤⓧ', latex_str)
        latex_str = re.sub(r'(∫|∑)_([a-zA-Z0-9])\^([a-zA-Z0-9])', r'ⓍⒺ\1ⓔⓈ\2ⓢⓊ\3ⓤⓧ', latex_str)
        
        # Chỉ số trên và dưới ĐỒNG THỜI (Hóa học/Vật lý lồng ghép như SO_4^{2-})
        latex_str = re.sub(r'([a-zA-Z0-9_>\]/()ⒻⓕⓃⓝⒹⓓⓆⓠⒺⓔⒷⓑⓅⓟⓈⓢⓊⓤⓍⓧⓎⓨⓋⓥⓂⓜⓇⓡ]+)_\{([^{}]+)\}\^\{([^{}]+)\}', r'ⓏⒺ\1ⓔⓈ\2ⓢⓊ\3ⓤⓩ', latex_str)
        latex_str = re.sub(r'([a-zA-Z0-9_>\]/()ⒻⓕⓃⓝⒹⓓⓆⓠⒺⓔⒷⓑⓅⓟⓈⓢⓊⓤⓍⓧⓎⓨⓋⓥⓂⓜⓇⓡ]+)\^\{([^{}]+)\}_\{([^{}]+)\}', r'ⓏⒺ\1ⓔⓈ\3ⓢⓊ\2ⓤⓩ', latex_str)

        # Chỉ số dưới (\_ )
        latex_str = re.sub(r'([a-zA-Z0-9_>\]/()ⒻⓕⓃⓝⒹⓓⓆⓠⒺⓔⒷⓑⓅⓟⓈⓢⓊⓤⓍⓧⓎⓨⓋⓥⓂⓜⓇⓡ]+)_\{([^{}]+)\}', r'ⒷⒺ\1ⓔⓈ\2ⓢⓑ', latex_str)
        latex_str = re.sub(r'([a-zA-Z0-9_>\]/()ⒻⓕⓃⓝⒹⓓⓆⓠⒺⓔⒷⓑⓅⓟⓈⓢⓊⓤⓍⓧⓎⓨⓋⓥⓂⓜⓇⓡ]+)_([a-zA-Z0-9])', r'ⒷⒺ\1ⓔⓈ\2ⓢⓑ', latex_str)
        
        # Chỉ số trên (\^ )
        latex_str = re.sub(r'([a-zA-Z0-9_>\]/()ⒻⓕⓃⓝⒹⓓⓆⓠⒺⓔⒷⓑⓅⓟⓈⓢⓊⓤⓍⓧⓎⓨⓋⓥⓂⓜⓇⓡ]+)\^\{([^{}]+)\}', r'ⓅⒺ\1ⓔⓊ\2ⓤⓟ', latex_str)
        latex_str = re.sub(r'([a-zA-Z0-9_>\]/()ⒻⓕⓃⓝⒹⓓⓆⓠⒺⓔⒷⓑⓅⓟⓈⓢⓊⓤⓍⓧⓎⓨⓋⓥⓂⓜⓇⓡ]+)\^([a-zA-Z0-9])', r'ⓅⒺ\1ⓔⓊ\2ⓤⓟ', latex_str)

        if latex_str == prev: break

    # 5. Render thành chuẩn OMML XML
    MARKERS = "".join(TAG_MAP.keys())
    parts = re.split(f'([{MARKERS}])', latex_str)
    
    RPR = '<m:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:sz w:val="28"/><w:szCs w:val="28"/></m:rPr>'
    omml_xml = f'<m:oMath {nsdecls("m")}>'
    
    for p in parts:
        if p in TAG_MAP:
            omml_xml += TAG_MAP[p]
        elif p.strip() or p == ' ':
            omml_xml += f'<m:r>{RPR}<m:t>{p}</m:t></m:r>'
            
    omml_xml += '</m:oMath>'
    
    try:
        return parse_xml(omml_xml)
    except Exception as e:
        return None

# ================= BÓC TÁCH MỌI KIỂU TOÁN HỌC =================
def process_runs_with_math(paragraph, text):
    """Tự động nhận diện mọi kiểu định dạng Toán: $$ $$, \[ \], \( \), $ $"""
    text_clean = text.strip()
    
    # Chia tách thông minh mọi loại thẻ toán học
    delimiter_pattern = r'(\$\$[\s\S]*?\$\$|\\\[[\s\S]*?\\\]|\\\([\s\S]*?\\\)|\$[\s\S]*?\$)'
    parts = re.split(delimiter_pattern, text_clean)
    
    for part in parts:
        if not part: continue
        
        # Xác định nội dung bên trong cặp thẻ
        is_math = False
        math_content = ""
        if part.startswith('$$') and part.endswith('$$'):
            math_content = part[2:-2].strip(); is_math = True
        elif part.startswith('\\[') and part.endswith('\\]'):
            math_content = part[2:-2].strip(); is_math = True
        elif part.startswith('\\(') and part.endswith('\\)'):
            math_content = part[2:-2].strip(); is_math = True
        elif part.startswith('$') and part.endswith('$'):
            math_content = part[1:-1].strip(); is_math = True

        # Render OMML nếu là toán học, hoặc in thường nếu là văn bản
        if is_math and math_content:
            math_element = convert_latex_to_omml(math_content)
            if math_element is not None:
                paragraph._p.append(math_element)
            else:
                run = paragraph.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        elif not is_math:
            bold_parts = part.split('**')
            for i, b_part in enumerate(bold_parts):
                is_bold = (i % 2 != 0)
                sub_sup_parts = re.split(r'(<sub>.*?</sub>|<sup>.*?</sup>)', b_part)
                for s_part in sub_sup_parts:
                    if not s_part: continue
                    if s_part.startswith('<sub>') and s_part.endswith('</sub>'):
                        run = paragraph.add_run(s_part[5:-6])
                        run.bold = is_bold
                        run.font.subscript = True
                        run.font.name, run.font.size = 'Times New Roman', Pt(14)
                    elif s_part.startswith('<sup>') and s_part.endswith('</sup>'):
                        run = paragraph.add_run(s_part[5:-6])
                        run.bold = is_bold
                        run.font.superscript = True
                        run.font.name, run.font.size = 'Times New Roman', Pt(14)
                    else:
                        run = paragraph.add_run(s_part)
                        run.bold = is_bold
                        run.font.name, run.font.size = 'Times New Roman', Pt(14)

# ================= XỬ LÝ ĐỒ THỊ AI =================
def generate_plot_stream(eq_str):
    fig, ax = plt.subplots(figsize=(5, 3.5))
    x = np.linspace(-10, 10, 400)
    safe_dict = {"x": x, "np": np, "sin": np.sin, "cos": np.cos, "tan": np.tan, "sqrt": np.sqrt}
    try:
        eq_str_py = eq_str.replace('^', '**')
        y = eval(eq_str_py, {"__builtins__": {}}, safe_dict)
        if isinstance(y, (int, float)):
            y = np.full_like(x, y)
        ax.plot(x, y, color='#1E40AF', linewidth=2.5)
        ax.axhline(0, color='black', linewidth=1.2)
        ax.axvline(0, color='black', linewidth=1.2)
        ax.grid(color='gray', linestyle='--', linewidth=0.5, alpha=0.7)
        ax.set_ylim([-10, 10])
        ax.set_title(f"Đồ thị: y = {eq_str}", fontsize=10, pad=10)
    except:
        ax.text(0.5, 0.5, f"[Lỗi cú pháp vẽ đồ thị]", ha='center', va='center', color='red')
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf
