import streamlit as st
import io
import re
from docx import Document
from google import genai
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime

# ================= CẤU HÌNH GOOGLE SHEETS =================
SHEET_ID = '1C6642jk_oQ0g9UC2By2qsNxxfQVR0MrZYj52tRdWDlY'

def get_sheet():
    creds_dict = dict(st.secrets["GOOGLE_KEY"])
    scopes = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
    creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)
    client = gspread.authorize(creds)
    return client.open_by_key(SHEET_ID).sheet1

# ================= HÀM HỖ TRỢ =================
def create_word_file(title, content):
    doc = Document()
    doc.add_heading(title, 0)
    clean_content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE) 
    clean_content = clean_content.replace('**', '') 
    clean_content = re.sub(r'^\*\s+', '- ', clean_content, flags=re.MULTILINE) 
    clean_content = clean_content.replace('*', '') 
    doc.add_paragraph(clean_content)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ================= CÁC THẺ GIAO DIỆN =================
def render_tab_1():
    st.info("💡 **THẺ 1 - BỘ LỌC TÌM KIẾM:** Chọn các tiêu chí để AI gợi ý dự án STEM.")
    col_m1, col_m2, col_m3 = st.columns(3)
    with col_m1: chon_khoi_t1 = st.selectbox("📌 Khối lớp:", ["Khối 9", "Khối 8", "Khối 7", "Khối 6"], key="khoi_t1")
    with col_m2: chon_mon_t1 = st.selectbox("📌 Môn chủ đạo:", ["Khoa học tự nhiên", "Toán học", "Công nghệ", "Tin học"], key="mon_t1")
    with col_m3: chon_chu_de_t1 = st.selectbox("📌 Lĩnh vực:", ["Tiết kiệm năng lượng", "Bảo vệ môi trường", "Nông nghiệp thông minh", "Nhà thông minh", "Chủ đề tự do"], key="chude_t1")
    
    mon_tich_hop_t1 = st.multiselect("➕ Các môn học cần tích hợp:", ["Toán học", "Khoa học", "Công nghệ", "Kỹ thuật", "Nghệ thuật"], key="montichhop_t1")
    tich_hop_ai_t1 = st.checkbox("🔌 Yêu cầu AI ưu tiên dự án có Vi điều khiển", value=True, key="chk_ai_t1")
    tich_hop_khuyet_tat_t1 = st.checkbox("🤝 Ưu tiên dự án phù hợp Giáo dục hòa nhập", value=True, key="chk_kt_t1")
    
    if st.button("✨ KÍCH HOẠT AI GỢI Ý CHỦ ĐỀ MỚI", use_container_width=True, key="btn_ai_t1"):
        with st.spinner("AI đang tổng hợp dữ liệu..."):
            prompt_goi_y = f"Đề xuất 3 dự án STEM cho: {chon_khoi_t1}, Môn: {chon_mon_t1}, Lĩnh vực: {chon_chu_de_t1}. Yêu cầu: {tich_hop_ai_t1 and 'Có vi điều khiển' or 'Không bắt buộc'}."
            client = genai.Client() 
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt_goi_y)
            st.session_state.stem_ai_suggestions = response.text
            st.rerun()
    
    if st.session_state.stem_ai_suggestions:
        with st.container(border=True):
            st.markdown(st.session_state.stem_ai_suggestions)

def render_tab_2():
    st.success("🛠️ **THẺ 2 - THIẾT KẾ KHBD:** Nhập tên dự án để AI soạn giáo án.")
    ten_chu_de_t2 = st.text_input("✍️ Nhập Tên dự án:", key="ten_t2")
    col1, col2 = st.columns(2)
    with col1:
        mon_chu_dao_t2 = st.selectbox("📚 Môn học:", ["Khoa học tự nhiên", "Toán học", "Công nghệ", "Tin học"], key="mon_t2")
        lop_hoc_t2 = st.selectbox("🎓 Dành cho:", ["Lớp 9", "Lớp 8", "Lớp 7", "Lớp 6"], key="lop_t2")
    with col2:
        thoi_luong_t2 = st.text_input("⏱️ Thời lượng:", key="thoiluong_t2")
    
    tich_hop_ai_iot_t2 = st.checkbox("🔌 Tích hợp Vi điều khiển", value=True, key="chk_ai_t2")
    
    if st.button("🚀 KÍCH HOẠT AI BIÊN SOẠN KHBD CHI TIẾT", type="primary", use_container_width=True, key="btn_ai_t2"):
        with st.spinner("AI đang soạn KHBD..."):
            prompt = f"Soạn KHBD STEM chi tiết cho: {ten_chu_de_t2}. Môn: {mon_chu_dao_t2}. Tích hợp IoT: {tich_hop_ai_t2}."
            client = genai.Client()
            response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt)
            st.session_state.stem_generated_content = response.text
            st.rerun()

    if st.session_state.stem_generated_content:
        st.markdown(st.session_state.stem_generated_content)
        if st.button("💾 LƯU VÀO GOOGLE SHEETS"):
            get_sheet().append_row([ten_chu_de_t2, st.session_state.stem_generated_content, datetime.now().strftime("%Y-%m-%d")])
            st.toast("Đã lưu thành công lên Sheets!", icon="✅")

def render_tab_3():
    st.warning("📁 **THẺ 3 - KHO LƯU TRỮ:**")
    st.info("Dữ liệu của thầy đã được đồng bộ an toàn lên Google Sheets.")

def render_stem_section():
    st.markdown("## 🚀 HỆ SINH THÁI GIÁO DỤC STEM")
    if "stem_generated_content" not in st.session_state: st.session_state.stem_generated_content = ""
    if "stem_ai_suggestions" not in st.session_state: st.session_state.stem_ai_suggestions = ""
    
    tab1, tab2, tab3 = st.tabs(["💡 1. SẢN PHẨM", "🛠️ 2. XÂY DỰNG KHBD", "📁 3. KHBD ĐÃ LƯU"])
    with tab1: render_tab_1()
    with tab2: render_tab_2()
    with tab3: render_tab_3()
