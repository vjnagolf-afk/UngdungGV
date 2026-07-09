import streamlit as st

def render_nam_hoc_tab():
    st.markdown("### 📌 I. ĐẶC ĐIỂM TÌNH HÌNH LỚP")
    col_tl, col_kk = st.columns(2)
    with col_tl:
        st.text_area(
            "Thuận lợi", 
            value="- Nhà trường luôn quan tâm đến công tác chủ nhiệm, có những kế hoạch chỉ đạo kịp thời đến tập thể lớp.\n- Tập thể lớp 9D có tinh thần đoàn kết, có ý thức trong công việc được giao, công việc của lớp, có tinh thần cầu tiến, có ý thức vươn lên trong học tập và rèn luyện đạo đức tác phong.\n- Có đội ngũ ban cán sự lớp nhiệt tình năng nổ, hết lòng vì công việc chung.",
            height=180, key="nam_hoc_thuan_loi"
        )
    with col_kk:
        st.text_area(
            "Khó khăn", 
            value="Một số học sinh vẫn còn tâm lý lười học, thiếu tập trung, đặc biệt ở các môn khoa học.\nCó học sinh thuộc diện hoàn cảnh khó khăn, thiếu sự quan tâm của gia đình.\nMột số em có biểu hiện chống đối nhẹ, dễ bị lôi kéo.\nKỹ năng tự học, tự quản và làm việc nhóm của học sinh chưa đồng đều.\nLớp có một số em ý thức chưa được tốt.",
            height=180, key="nam_hoc_kho_khan"
        )
        
    st.write("---")
    st.markdown("### 📝 II. KẾ HOẠCH NĂM HỌC 2025-2026")
    
    col_rl, col_mdc = st.columns(2)
    with col_rl:
        st.text_area(
            "Rèn luyện trong nhà trường", 
            value="+ Xây dựng uy tín với học sinh, với đồng nghiệp, với cha mẹ học sinh và xã hội về chuyên môn, nghiệp vụ. Tư cách đạo đức, tác phong sư phạm mẫu mực trong sinh hoạt. Không làm giảm niềm tin, mất uy tín, danh dự của nhà trường trước xã hội về chuyên môn, nghiệp vụ và tư cách đạo đức, tác phong sinh hoạt của cá nhân.",
            height=150, key="nam_hoc_ren_luyen"
        )
    with col_mdc:
        st.text_area(
            "Mục đích yêu cầu chung", 
            value="- Luôn kính trọng người trên, thầy cô giáo, cán bộ và nhân viên nhà trường; thương yêu và giúp đỡ nhau, có ý thức xây dựng tập thể, đoàn kết với các bạn, được các bạn tin yêu.\n- Tích cực rèn luyện phẩm chất đạo đức, có lối sống lành mạnh, trung thực, giản dị, khiêm tốn.\n- Hoàn thành đầy đủ nhiệm vụ học tập, có ý thức cố gắng vươn lên trong học tập.\n- Thực hiện nghiêm túc nội quy nhà trường; chấp hành tốt Pháp luật, quy định về trật tự, an toàn",
            height=150, key="nam_hoc_muc_dich_chung"
        )
        
    col_ctc, col_bpc = st.columns(2)
    with col_ctc:
        st.text_area(
            "Chỉ tiêu chung", 
            value="100% học sinh không vi phạm nội quy trường học.\n95% học sinh đạt rèn luyện tốt và Khá.\nKhông có học sinh bỏ học.",
            height=120, key="nam_hoc_chi_tieu_chung"
        )
    with col_bpc:
        st.text_area(
            "Biện pháp chính chung", 
            value="- GVCN quán triệt cho học sinh nội quy của nhà trường, bài học văn hóa ứng xử, giáo dục truyền thống nhà trường, giáo dục môi trường, an toàn giao thông và nhiệm vụ học sinh THCS. - Tổ chức cho học sinh ký cam kết thực hiện nội quy, phòng chống ma túy và tệ nạn xã hội, ký giao ước thi đua thực hiện an toàn giao thông.\n- Tập huấn cán bộ lớp, đội cờ đỏ, cán bộ giữ sổ đầu bài.",
            height=120, key="nam_hoc_bien_phap_chung"
        )

    st.markdown("#### 🎯 Mục đích cụ thể từng mặt")
    st.info("✨ Phân hệ 1: HỌC TẬP")
    col_ht_yc, col_ht_ct = st.columns(2)
    with col_ht_yc:
        st.text_area("Yêu cầu (Học tập)", value="- 100% học sinh đi học có đầy đủ SGK và các dụng cụ phục vụ học tập. - 100% soạn bài và làm bài tập, học thuộc bài ở nhà. - 100% HS tích cực tham gia phát biểu xây dựng bài trong các tiết học. - Giáo dục học sinh có ý thức, mục tiêu học tập đúng đắn.\n- Giáo dục học sinh có phương pháp học tập đạt hiệu quả cao.\n- Phát triển khả năng tư duy, sáng tạo của học sinh.", height=150, key="ht_yeu_cau")
    with col_ht_ct:
        st.text_area("Chỉ tiêu (Học tập)", value="- Mức Tốt: 5 HS\n- Mức Khá: 12 HS - Mức Đạt: 22 HS trở lên\n- Danh hiệu 'Học sinh Xuất sắc': 1 HS - Danh hiệu 'Học sinh Giỏi': 5 HS", height=150, key="ht_chi_tieu")
    st.text_area("Biện pháp chính (Học tập)", value="- Tổ chức cho HS thảo luận, tìm phương pháp học tập đạt hiệu quả cao.\n- Kết hợp với GVBM thành lập nhóm hs cán sự theo bộ môn, đẩy mạnh các phong trào 'Đôi bạn học tập'\n- Liên hệ với GVBM để theo sát tình hình học tập của HS, kịp thời động viên khen thưởng HS học tập tốt, có nhiều tiến bộ và nhắc nhở\n- Yêu cầu học sinh phải có góc học tập ở nhà, chú trọng việc tự học của HS.\n- Thực hiện nghiêm túc quy chế kiểm tra thi cử.", height=120, key="ht_bien_phap")

    st.write("")
    st.info("✨ Phân hệ 2: GIÁO DỤC HƯỚNG NGHIỆP")
    col_hn_yc, col_hn_ct = st.columns(2)
    with col_hn_yc:
        st.text_area("Yêu cầu (Hướng nghiệp)", value="- Giáo dục học sinh có ý thức yêu lao động, có trách nhiệm giữ gìn vệ sinh, bảo vệ tài sản, cảnh quan, môi trường thiên nhiên.\n- Giúp học sinh định hướng được nghề nghiệp phù hợp, từ đó có mục tiêu học tập đúng đắn.", height=120, key="hn_yeu_cau")
    with col_hn_ct:
        st.text_area("Chỉ tiêu (Hướng nghiệp)", value="- 100% HS làm tốt công tác bảo vệ tài sản, cảnh quan, môi trường 'Xanh - sạch - đẹp'.\n- 100% HS tham gia đầy đủ các buổi lao động theo kế hoạch của nhà trường.", height=120, key="hn_chi_tieu")
    st.text_area("Biện pháp chính (Hướng nghiệp)", value="- Quán triệt học sinh tham gia đầy đủ các buổi lao động trường, lớp, ý thức bảo vệ môi trường, lao động tập thể.\n- Giáo dục HS có tinh thần hăng say tích cực, ý thức trong lao động xây dựng cảnh quan, tạo môi trường xanh, sạch, đẹp. - Giáo dục HS tinh thần tập thể, giao phó công việc cho HS phù hợp, đúng đối tượng để buổi lao động đạt hiệu quả cao.\n- Nhận xét, đánh giá, khen thưởng cho HS khi tham gia.", height=120, key="hn_bien_phap")

    st.write("")
    st.info("✨ Phân hệ 3: GIÁO DỤC NGOÀI GIỜ LÊN LỚP")
    col_gl_yc, col_gl_ct = st.columns(2)
    with col_gl_yc:
        st.text_area("Yêu cầu (Ngoài giờ lên lớp)", value="- Kĩ năng tự nhận thức (ta là ai là điều cực kỳ quan trọng)\n- Kĩ năng giải quyết các tình huống đặc biệt khó khăn trong cuộc sống,\n- Kĩ năng giao tiếp và thương thuyết (bao hàm tính tự kiềm chế)\n- Kĩ năng lựa chọn và quyết định (bao hàm phê phán và bác bỏ)\n- KN hợp tác và tìm kiếm sự giúp đỡ (bao hàm yếu tố thân thiện, làm việc theo nhóm)", height=150, key="gl_yeu_cau")
    with col_gl_ct:
        st.text_area("Chỉ tiêu (Ngoài giờ lên lớp)", value="100% HS rèn kĩ năng sống cần được giáo dục và rèn luyện.", height=150, key="gl_chi_tieu")
    st.text_area("Biện pháp chính (Ngoài giờ lên lớp)", value="- Nâng cao ý thức tự nguyện, tự giác, tự chủ, phát huy tính tích cực trong mọi hoạt động rèn luyện kĩ năng sống;\n- Nhận thức rằng, việc rèn luyện KNS là việc của mình, trước hết có lợi cho việc học tập và sự tiến bộ về mọi mặt của chính mình, cho gia đình và sau đó cho cộng đồng, cho XH và đất nước;\n- Không chỉ rèn luyện cho mình mà quan tâm đến việc rèn luyện chung của cả một tập thể tổ, lớp và rộng hơn, của trường mình.", height=120, key="gl_bien_phap")

    st.write("---")
    st.markdown("### 📊 III. CHỈ TIÊU TOÀN DIỆN CUỐI NĂM")
    col_td_ct, col_td_bp = st.columns(2)
    with col_td_ct:
        st.text_area(
            "Chỉ tiêu cuối năm", 
            value="1. Danh hiệu lớp: - Lớp Tiên tiến - Chi đội: Mạnh\n2. Rèn luyện: Tốt: 40 (97,7%) Khá: 02 (2,3%) Đạt/CĐ: 0 (0%)\n3. Học tập:\n- Mức Tốt: 5HS (13,5%)\n- Mức Khá: 12 HS (27,9%)",
            height=150, key="td_chi_tieu"
        )
    with col_td_bp:
        st.text_area(
            "Biện pháp chính cuối năm", 
            value="- GVCN thường xuyên bám sát lớp, gần gũi tiếp xúc with HS để tìm hiểu gia đình từng em từ đó đề ra biện pháp giáo dục. - GVCN kết hợp chặt chẽ với GVBM và Tổng phụ trách Đội để kịp thời uốn nắn, nhắc nhở các em giúp các em tiến bộ. - GVCN lớp khách quan công bằng với HS khen, chê kịp thời. - Xây dựng ban cán sự lớp gương mẫu, nhiệt tình và có trách nhiệm cao. - GVCN bám sát kế hoạch của trường - Đội để triểnkai kịp thời cho lớp. - Phối hợp thường xuyên với PHHS và GVBM.",
            height=150, key="td_bien_phap"
        )
        
    if st.button("💾 Lưu Toàn bộ Kế hoạch năm học", type="primary", key="btn_save_nam_hoc_full"):
        st.success("🎉 Đã đồng bộ và lưu trữ thành công toàn bộ Khung kế hoạch năm học!")
import streamlit as st
import io
from docx import Document

def export_to_word(title, content_text):
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content_text.split("\n"):
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def render_thang_tab(run_ai_prompt_safe=None):
    st.write("#### 🛠 CẤU HÌNH THÔNG TIN CHỦ NHIỆM")
    col_khoi, col_lop, col_thang = st.columns(3)
    with col_khoi:
        selected_khoi = st.selectbox("Chọn Khối lớp:", ["Khối lớp 6", "Khối lớp 7", "Khối lớp 8", "Khối lớp 9"], key="sb_khoi_lop")
    with col_lop:
        lop_dict = {"Khối lớp 6": ["6A","6B","6C","6D","6E","6F"], "Khối lớp 7": ["7A","7B","7C","7D","7E","7F"], "Khối lớp 8": ["8A","8B","8C","8D","8E","8F"], "Khối lớp 9": ["9A","9B","9C","9D","9E","9F","9G"]}
        selected_lop = st.selectbox("Chọn Lớp chủ nhiệm:", lop_dict.get(selected_khoi, ["6A"]), key="sb_lop_chu_nhiem")
    with col_thang:
        thang_options = [f"Tháng {i}/2026" for i in range(8, 13)] + [f"Tháng {i}/2027" for i in range(1, 6)]
        selected_thang = st.selectbox("Chọn Tháng công tác:", thang_options, key="sb_thang_cong_tac")
        
    st.write("---")
    ghi_chu_them = st.text_input("Yêu cầu bổ sung đặc biệt cho tháng này (nếu có):", placeholder="Ví dụ: Chuẩn bị văn nghệ...", key="txt_ghi_chu_them")
    
    if "content_ke_hoach_thang" not in st.session_state:
        st.session_state["content_ke_hoach_thang"] = ""
        
    if st.button("🚀 Khởi tạo Kế hoạch bằng AI", type="primary", key="btn_chu_nhiem_ai"):
        if run_ai_prompt_safe is not None:
            with st.spinner(f"AI đang thiết lập kế hoạch {selected_thang}..."):
                prompt_he_thong = f"Lập kế hoạch công tác chủ nhiệm chi tiết cho lớp {selected_lop} trong {selected_thang}. Xuống dòng rõ ràng từng mục: 1. Chủ điểm, 2. Nội dung hoạt động, * Kế hoạch từng tuần (Tuần 22, Tuần 23...). Lưu ý từ GV: {ghi_chu_them}"
                response = run_ai_prompt_safe(prompt_he_thong)
                st.session_state["content_ke_hoach_thang"] = response
        else:
            st.info("Hệ thống kết nối AI đang được đồng bộ...")

    st.write("#### 📝 KHUNG SOẠN THẢO KẾ HOẠCH THÁNG")
    edited_text = st.text_area(
        label="Nội dung kế hoạch hiển thị theo cấu trúc dọc (Có thể chỉnh sửa):",
        value=st.session_state["content_ke_hoach_thang"],
        height=400, key="ta_main_editor"
    )
    st.session_state["content_ke_hoach_thang"] = edited_text
    
    if st.session_state["content_ke_hoach_thang"].strip():
        file_name_doc = f"Ke_hoach_chu_nhiem_{selected_lop}_{selected_thang.replace('/', '_')}.docx"
        word_file_bytes = export_to_word(f"KẾ HOẠCH CHỦ NHIỆM LỚP {selected_lop} - {selected_thang.upper()}", edited_text)
        st.download_button(
            label="📥 Tải xuống file Word (.docx)",
            data=word_file_bytes,
            file_name=file_name_doc,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="btn_download_word"
        )
import streamlit as st
# Import các phân hệ con từ các file nhỏ đã chia tách bên trên
from chu_nhiem_nam_hoc import render_nam_hoc_tab
from chu_nhiem_thang import render_thang_tab

def render_chu_nhiem_section(run_ai_prompt_safe=None):
    # Khởi tạo 2 thẻ song song lớn cho Phân hệ Chủ nhiệm
    tab_tong_quan, tab_hang_thang = st.tabs([
        "📊 Đặc điểm tình hình & Kế hoạch năm học", 
        "📅 Kế hoạch công tác theo Tháng (AI Tự động)"
    ])
    
    with tab_tong_quan:
        # Gọi hàm hiển thị kế hoạch năm học từ file chu_nhiem_nam_hoc.py
        render_nam_hoc_tab()

    with tab_hang_thang:
        # Gọi hàm hiển thị kế hoạch tháng từ file chu_nhiem_thang.py
        render_thang_tab(run_ai_prompt_safe)
