import streamlit as st
import docx  
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
from pypdf import PdfReader
import gspread
from datetime import datetime

from math_compiler import process_runs_with_math, generate_plot_stream

# ================= CẤU HÌNH GOOGLE SHEETS =================
SHEET_ID = '1C6642jk_oQ0g9UC2By2qsNxxfQVR0MrZYj52tRdWDlY' 

def get_khbd_sheet():
    try:
        creds_dict = None
        priority_keys = ["gspread_credentials", "GSPREAD_CREDENTIALS", "google_sheet_creds", "google_creds", "GOOGLE_KEY"]
        for key in priority_keys:
            if key in st.secrets:
                creds_dict = st.secrets[key]
                break
                
        if creds_dict is None:
            for key in st.secrets.keys():
                node = st.secrets[key]
                if hasattr(node, "get") or isinstance(node, dict):
                    if node.get("type") == "service_account" or "private_key" in node:
                        creds_dict = node
                        break
                        
        if creds_dict is None: return None
        gc = gspread.service_account_from_dict(creds_dict)
        sh = gc.open_by_key(SHEET_ID)
        return sh.worksheet("KHBD")
    except:
        return None

# --- HÀM TRÍCH XUẤT VĂN BẢN VÀ LỌC ẢNH TRÙNG LẶP ---
def extract_context_from_uploaded_files(uploaded_files):
    combined_text = ""
    extracted_images = [] 
    seen_image_sizes = set()
    for file in uploaded_files:
        try:
            if file.name.endswith('.docx'):
                doc = docx.Document(file)
                for paragraph in doc.paragraphs:
                    if paragraph.text: combined_text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        text_row = [cell.text for cell in row.cells]
                        combined_text += " | ".join(text_row) + "\n"
                for rel in doc.part.relations.values():
                    if "image" in rel.target_ref:
                        img_blob = rel.target_part.blob
                        if len(img_blob) not in seen_image_sizes:
                            seen_image_sizes.add(len(img_blob))
                            extracted_images.append(img_blob)
            elif file.name.endswith('.pdf'):
                reader = PdfReader(file)
                for page in reader.pages:
                    combined_text += (page.extract_text() or "") + "\n"
                    for img_file_object in page.images:
                        img_blob = img_file_object.data
                        if len(img_blob) not in seen_image_sizes:
                            seen_image_sizes.add(len(img_blob))
                            extracted_images.append(img_blob)
            elif file.name.endswith('.txt'):
                combined_text += file.read().decode("utf-8") + "\n"
        except Exception as e:
            st.error(f"Lỗi khi xử lý file {file.name}: {str(e)}")
    return combined_text, extracted_images

def set_paragraph_spacing(paragraph, before_pt=3.0, after_pt=4.5):
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    p_pr.append(spacing)

# ================= ĐỘNG CƠ KẾT XUẤT ĐỊNH DẠNG CỨNG =================
def export_khbd_to_docx(markdown_content, images_list):
    # LÀM SẠCH GỐC: Loại bỏ hoàn toàn định dạng Markdown thừa
    markdown_content = re.sub(r'(?m)^#+\s*', '', markdown_content)
    # Loại bỏ các dòng định dạng bảng markdown thừa |---|
    markdown_content = re.sub(r'\|[:\-\s]+\|', '', markdown_content)
    
    doc = docx.Document()
    # THIẾT LẬP LỀ 
    for section in doc.sections:
        section.top_margin = Inches(0.79); section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18); section.right_margin = Inches(0.59)

    MAU_DO = RGBColor(255, 0, 0); MAU_XANH_DUONG = RGBColor(0, 51, 153); MAU_DEN = RGBColor(0, 0, 0)

    lines = markdown_content.split('\n')
    
    # BIẾN LƯU BẢNG TẠM
    table_rows = []
    
    for line in lines:
        # HỦY MỌI ĐỊNH DẠNG AI (Markdown) để lấy văn bản thuần
        cl = line.strip().replace('**', '').replace('*', '').replace('#', '')
        if not cl: continue

        # XỬ LÝ BẢNG (Markdown Table)
        if '|' in cl and not cl.startswith(('-', 'I.', 'II.', '1.', '2.')):
            row_data = [cell.strip() for cell in cl.split('|') if cell.strip()]
            if row_data: table_rows.append(row_data)
            continue
        else:
            if table_rows:
                max_cols = max(len(row) for row in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                tbl.style = 'Table Grid'
                
                for i, row in enumerate(table_rows):
                    for j in range(max_cols):
                        val = row[j] if j < len(row) else ""
                        # [BẢN VÁ]: KHÔNG DÙNG cell.text, SỬ DỤNG PARAGRAPH ĐỂ RENDER TOÁN
                        cell = tbl.cell(i, j)
                        # Xóa paragraph mặc định của cell
                        for p in cell.paragraphs:
                            p = p._element
                            p.getparent().remove(p)
                        # Thêm paragraph mới và dùng trình biên dịch toán
                        p_cell = cell.add_paragraph()
                        process_runs_with_math(p_cell, val)
                        # Định dạng font cho ô bảng
                        for r in p_cell.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(12)
                
                table_rows = []
                doc.add_paragraph()

        # XỬ LÝ ĐỀ MỤC VÀ ĐỊNH DẠNG (KHÓA CỨNG)
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 3.0, 4.5)
        p.paragraph_format.left_indent = Inches(0); p.paragraph_format.right_indent = Inches(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Định dạng theo Rule:
        upper_cl = cl.upper()
        
        # RULE 1: TIÊU ĐỀ BÀI (Đỏ, In hoa, Giữa)
        if upper_cl.startswith(("KẾ HOẠCH BÀI DẠY", "BÀI")):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(upper_cl)
            run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = MAU_DO
        
        # RULE 2: MÔN, LỚP, TIẾT (Xanh, Đậm, Giữa)
        elif any(upper_cl.startswith(x) for x in ["MÔN HỌC", "LỚP", "THỜI LƯỢNG", "TIẾT"]):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(upper_cl)
            run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = MAU_XANH_DUONG
            
        # RULE 3: ĐỀ MỤC LỚN (I. II. III.) (Xanh, Đậm, Hoa)
        elif re.match(r'^(I|II|III|IV|V|VI|VII)\.', cl):
            process_runs_with_math(p, upper_cl)
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.font.color.rgb = MAU_XANH_DUONG
            
        # RULE 4: ĐỀ MỤC NHỎ (1. Kiến thức, 2. Năng lực...) (Đỏ, Đậm)
        elif re.match(r'^\d+\.\s+(Kiến thức|Năng lực|Phẩm chất)', cl, re.IGNORECASE):
            process_runs_with_math(p, cl)
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.font.color.rgb = MAU_DO
            
        # RULE 5: TIỀN TỐ ĐẬM (a) Mục tiêu, Bước 1, 2.1 Hoạt động...)
        else:
            m = re.match(r'^([a-d]\)\s*[^:]+:?|\d+\.\d+\.?\s*Hoạt động\s*\d*\.?\s*:?|Bước\s+\d+:?)(.*)', cl, re.IGNORECASE)
            if m:
                r_pref = p.add_run(m.group(1).strip() + " ")
                r_pref.bold = True; r_pref.font.name = 'Times New Roman'; r_pref.font.size = Pt(14); r_pref.font.color.rgb = MAU_DEN
                if m.group(2): process_runs_with_math(p, m.group(2).strip())
            else:
                process_runs_with_math(p, cl)
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.font.color.rgb = MAU_DEN

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# (Giữ nguyên các phần import và hàm get_khbd_sheet, extract_context_from_uploaded_files, set_paragraph_spacing, export_khbd_to_docx như bản trước)

def render_khbd_section(run_ai_prompt_safe_func):
    st.markdown("<h3 style='text-align: center; color: blue;'>🧠 TRỢ LÝ THIẾT KẾ KẾ HOẠCH BÀI DẠY (KHBD) AI</h3>", unsafe_allow_html=True)
    
    # Khởi tạo trạng thái
    if "ket_qua_khbd" not in st.session_state: st.session_state["ket_qua_khbd"] = ""
    
    tab_thiet_ke, tab_thu_vien = st.tabs(["📝 THIẾT KẾ KHBD TỰ ĐỘNG", "🗄️ THƯ VIỆN BÀI SOẠN"])
    
    with tab_thiet_ke:
        ten_bai = st.text_input("Tên bài học / Chủ đề bài dạy:", placeholder="Ví dụ: Bài 4: Tốc độ chuyển động - Khoa học tự nhiên 7")
        
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("⚡ Thiết kế bài dạy bằng AI", type="primary", use_container_width=True):
                if ten_bai:
                    with st.spinner("Đang soạn giáo án..."):
                        ket_qua, _ = run_ai_prompt_safe_func(f"Soạn KHBD chuẩn 5512 bài: {ten_bai}")
                        st.session_state["ket_qua_khbd"] = ket_qua
                else:
                    st.warning("⚠️ Vui lòng điền tên bài học!")
        
        with col2:
            # NÚT LƯU TẠM THỜI
            if st.button("💾 Lưu tạm thời vào Google Sheets", use_container_width=True):
                if st.session_state["ket_qua_khbd"]:
                    if save_khbd_to_sheet(ten_bai or "Bài không tên", "Lớp 7", "Kết nối tri thức", "2 tiết", st.session_state["ket_qua_khbd"]):
                        st.success("✅ Đã lưu vào Sheet!")
                    else:
                        st.error("❌ Không thể lưu vào Sheet.")
                else:
                    st.warning("⚠️ Chưa có bài để lưu!")
        
        if st.session_state["ket_qua_khbd"]:
            st.markdown("### 📋 Kết quả xem trước:")
            st.markdown(st.session_state["ket_qua_khbd"])
            
            # Nút tải file
            docx_data = export_khbd_to_docx(st.session_state["ket_qua_khbd"], [])
            st.download_button("📥 Tải tệp Word (.docx)", data=docx_data, file_name="KHBD.docx", use_container_width=True)

    with tab_thu_vien:
        st.write("### 📂 Danh sách bài soạn đã lưu trên Google Sheets")
        ds_bai = get_all_khbd_from_sheet()
        
        if not ds_bai:
            st.info("Chưa có bài soạn nào được lưu.")
        else:
            for idx, bai in enumerate(ds_bai):
                with st.expander(f"📖 {bai.get('Tên bài', 'Bài soạn')} ({bai.get('Thời gian', 'N/A')})"):
                    col_a, col_b = st.columns(2)
                    # NÚT GỌI BÀI
                    if col_a.button("📥 Gọi lại bài soạn", key=f"load_{idx}", use_container_width=True):
                        st.session_state["ket_qua_khbd"] = bai['Nội dung chi tiết']
                        st.rerun()
                    # NÚT XÓA BÀI
                    if col_b.button("🗑️ Xóa bài soạn", key=f"del_{idx}", use_container_width=True):
                        if delete_khbd_from_sheet(idx):
                            st.success("✅ Đã xóa!")
                            st.rerun()
                        else:
                            st.error("❌ Xóa lỗi.")
